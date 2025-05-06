# Packages
import streamlit as st
import pandas as pd
from docx import Document

# Modules
import io
import time

# Utils
import lib

# The App config and styling
lib.app_config()

# Sidebar and Navbar
lib.sidebar()

st.title(lib.APP_TITLE)

st.divider()

# Steps
lib.steps(1)

st.divider()

st.header("Extract Raw Content")

st.markdown("<p>Extract paragraphs from a <b>CSV file</b> or from <b>Your Previous Backup</b> using this tool, getting raw contents.<br>As the names of the attendees may appear in the transcripts, it is possible to <b>hide them at this stage</b>.",unsafe_allow_html=True)

upload = False

if "saved_convert" not in st.session_state:
    st.subheader("What Do You Want To Extract")
    upload = st.file_uploader("**Please Upload your data as a csv file only** from previous tool ; unless be sure to have **Meeting, Topic, Item, Segment, Attendee and Message** in the first row of your csv file.", type=['csv'],accept_multiple_files=False)


if upload or "saved_convert" in st.session_state:
    error = ""
    if "saved_convert" in st.session_state:
        df = st.session_state["saved_convert"]["Data"]
    else:
        data = upload
        df = pd.read_csv(data,sep=None,engine='python')
        csv_error_string = "**There was an error processing this file !** Please be sure to upload only a csv file with: **Meeting, Topic, Item, Segment, Attendee and Message** in the first row. **Each column can't be empty** or have empty field(s)."

        good_header = set(['Meeting','Topic','Item','Segment','Attendee','Message']).issubset(df.columns)
        empty_columns = df.isnull().any()
        if not good_header or empty_columns.any():
            error = st.warning(csv_error_string, icon="⚠️")
                
    if error == "":
        st.write(f"**Reminder Of Your Data: {len(df)} rows** successfully loaded.")

        filter_informal = st.checkbox('**Clearing "informal messages" from the content**',key="filter_informal",value=True,disabled=True,help="This extraction process will automatically clear informal messages.")
        
        edited_data = df
        meeting = edited_data.iloc[0]["Meeting"]

        attendee_table = {
                "Attendee":     [],
                "Replacement":  []
            }
        with st.spinner('Analysing content... Please Wait.'):
            time.sleep(3)

        with st.spinner('Extracting content... Please Wait.'):
            topics_list = {}
            key_item = 0
            key_segment = 0
            key_attendee = 0
            key_message = 0
            
            attendee_key = 0
            for topic in edited_data["Topic"]:
                if topic and (topic not in topics_list):
                    topics_list[topic] = {}
            for item in edited_data["Item"]:
                if item and (item not in topics_list[edited_data.iloc[key_item]["Topic"]]) and item == edited_data.iloc[key_item]["Item"]:
                    topics_list[edited_data.iloc[key_item]["Topic"]][item] = {}
                key_item+=1
            for segment in edited_data["Segment"]:
                if segment and (segment not in topics_list[edited_data.iloc[key_segment]["Topic"]][edited_data.iloc[key_segment]["Item"]]) and segment == edited_data.iloc[key_segment]["Segment"]:
                    topics_list[edited_data.iloc[key_segment]["Topic"]][edited_data.iloc[key_segment]["Item"]][segment] = {}
                key_segment+=1
            for attendee in edited_data["Attendee"]:
                if attendee not in attendee_table["Attendee"]:
                    attendee_table["Attendee"].append(attendee)
                    attendee_table["Replacement"].append(f'Attendee{lib.add_zero(attendee_key)}')
                    attendee_key += 1
                if edited_data.iloc[key_attendee]["Segment"]:
                    topics_list[edited_data.iloc[key_attendee]["Topic"]][edited_data.iloc[key_attendee]["Item"]][edited_data.iloc[key_attendee]["Segment"]][attendee] = ""
                key_attendee+=1
            for message in edited_data["Message"]:
                if edited_data.iloc[key_message]["Segment"]: 
                    topics_list[edited_data.iloc[key_message]["Topic"]][edited_data.iloc[key_message]["Item"]][edited_data.iloc[key_message]["Segment"]][edited_data.iloc[key_message]["Attendee"]] += " " + message
                key_message+=1

            st.subheader("Hiding Speakers Names")

            st.write(f'You have **{len(attendee_table["Attendee"])} attendees** to hide.')

            attendee_table = st.data_editor(attendee_table,use_container_width=True)

            anonymize_attendee = st.checkbox('**Hide Speakers Names Now ?**',key="anonymize_attendee",value=True,help="This feature can hide automatically speakers names as they clearly appear in the transcripts.")

            st.write("""<div id='top-content'></div>""",unsafe_allow_html=True)

            st.subheader("Your Extracted Content's Below:")

            st.markdown("<p><a href='#download-or-save-your-data'>🔽Download Or Save Your Data At The Bottom Of This Page🔽</a></p>",unsafe_allow_html=True)

            content = st.container(border=True)

            content.write("""<div class='extracted-content' />""",unsafe_allow_html=True)

            topic_index = 0
            
            for topic in topics_list.keys():
                item_index = 0
                doc_download = Document()
                doc_txt = {}
                doc_txt[topic_index] = ""

                for item in topics_list[topic].keys():

                    count_chars = 0
                    
                    doc_download.add_heading(meeting,level=1)
                    doc_txt[topic_index] = doc_txt[topic_index] + meeting + "\n\n\n"
                    content.title(meeting)
                    count_chars = count_chars + len(meeting) + 3
                        
                    doc_download.add_heading(topic,level=2)
                    doc_txt[topic_index] = doc_txt[topic_index] + topic + "\n\n"
                    content.header( topic )
                    count_chars = count_chars + len(topic) + 2

                    doc_download.add_heading(item,level=3)
                    doc_txt[topic_index] = doc_txt[topic_index] + item + "\n"
                    content.subheader( item )
                    count_chars = count_chars + len( item ) + 1


                    for segment in topics_list[topic][item].keys():
                        breadcrumb = meeting + " > " + topic + " > " + item + " > " + segment
                        content.write( f"**{breadcrumb}**" )
                        doc_download.add_heading(breadcrumb,level=4)
                        doc_txt[topic_index] = doc_txt[topic_index] + breadcrumb + "\n"
                        count_chars = count_chars + len( breadcrumb ) + 1

                        for name,message in  topics_list[topic][item][segment].items():
                            
                            if anonymize_attendee:
                                index = attendee_table["Attendee"].index(name)
                                name = attendee_table["Replacement"][index]

                            if filter_informal:
                                message = lib.filter_informal(message)

                            if len(message.replace(" ","")) > 0:
                                message_clean = message.replace("$","\\$")
                                content.write( "**" + name + "**:" +message_clean,unsafe_allow_html=True)
                                doc_download.add_paragraph(name + ": " + message)
                                doc_txt[topic_index] = doc_txt[topic_index] + name + ": " + message + "\n"

                                count_chars = count_chars + len(name) + 2
                                count_chars = count_chars + len(message)
                
                    content.divider()
                    content.subheader("Download Or Save Your Data")
                    content.write(f"**{meeting} > {topic} > {item}**")
                    content.write(f"This content is {count_chars} chars long.")


                    if anonymize_attendee:
                        docx_filename = f'{lib.current_datetime}_{meeting}_{topic}_{item}_hidden_attendees_speeches.docx'
                    else:
                        docx_filename = f'{lib.current_datetime}_{meeting}_{topic}_{item}_originals_attendees_speeches.docx'
                    
                    col1, col2, col3 = content.columns(3)
                    with col1:
                        bio = io.BytesIO()
                        doc_download.save(bio)
                        if doc_download:
                            st.download_button(
                            label="Download Content As DOCX File",
                            type="secondary",
                            data=bio.getvalue(),
                            file_name=docx_filename,
                            mime="docx",
                            key=f"dowload_content_btn_{topic_index}_{item_index}"
                            )

                    with col2:
                        st.write("OR")
                    
                    with col3:
                        if st.button(label="**Save Data For Next Step >> Anonymize Content**",type="primary",key=f"save_content_btn_{topic_index}_{item_index}", on_click=lib.save_content,args=[f"{meeting}>{topic}>{item}",doc_txt[topic_index],attendee_table]):
                            st.switch_page("./pages/03_anonymize.py")
                        item_index = item_index + 1
                    
                    st.markdown("<p><a href='#top-content'>🔼Go Back To The Top Of This Content🔼</a></p>",unsafe_allow_html=True)
                topic_index = topic_index + 1              