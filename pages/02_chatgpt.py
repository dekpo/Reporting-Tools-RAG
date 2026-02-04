# Packages
import streamlit as st
import openai
from openai import OpenAI
from docx import Document
import pyperclip

# Modules
import time
import io
import os
from datetime import datetime

# Utils
import lib

# The App config and styling
lib.app_config()

# Sidebar and Navbar
lib.sidebar()

st.title(lib.APP_TITLE)

st.divider()

# Steps - updated to show position 2 in the new 4-step process
# lib.steps(1)

# st.divider()

if "gpt_api_key" not in st.session_state:
    st.header("ChatGPT Tool")

    st.markdown("<p>Ask ChatGPT to answer questions about your anonymized content using advanced retrieval techniques.</p>",unsafe_allow_html=True)

    # Check data sources and show unified status message
    persist_directory = './chroma_db'
    document_metadata = None
    if os.path.exists(persist_directory) and os.path.isdir(persist_directory):
        document_metadata = lib.load_document_metadata(persist_directory)
    
    tabular_metadata = lib.get_tabular_metadata()
    
    # Create unified message if we have any data sources
    if document_metadata or tabular_metadata:
        sources_info = []
        
        if document_metadata:
            active_docs = sum(1 for doc_data in document_metadata.values() if doc_data.get("active", True))
            total_docs = len(document_metadata)
            
            if active_docs == total_docs:
                sources_info.append(f"{total_docs} text document(s)")
            else:
                sources_info.append(f"{active_docs}/{total_docs} text document(s)")
        
        if tabular_metadata:
            active_datasets = sum(1 for dataset_data in tabular_metadata.values() if dataset_data.get("active", True))
            total_datasets = len(tabular_metadata)
            
            if active_datasets == total_datasets:
                sources_info.append(f"{total_datasets} tabular dataset(s)")
            else:
                sources_info.append(f"{active_datasets}/{total_datasets} tabular dataset(s)")
        
        if sources_info:
            sources_text = " and ".join(sources_info)
            st.info(f"Found {sources_text} in storage. After submitting your API key, you'll be able to access them.")
    
    # Check for existing data sources for the tip message
    has_documents = os.path.exists(persist_directory) and os.path.isdir(persist_directory) and lib.load_document_metadata(persist_directory)
    has_datasets = tabular_metadata is not None and len(tabular_metadata) > 0
    
    # Show message and button if no data sources are available
    if not has_documents and not has_datasets:
        st.info("💡 **Tip:** Upload documents or datasets on the Home page first to get the most out of ChatGPT analysis!")
        st.page_link(page="pages/00_home.py", label="Go To Home Page", icon=":material/home:", use_container_width=True)
        st.divider()

    MY_API_KEY = st.text_input(
            label="**Please specify your own OpenAI API Key** this one is for testing purpose you can use it and buy me a coffee ;)",
            value=st.secrets["OPENAI_API_KEY"]
        )

    st.markdown("<a href=\"https://platform.openai.com/api-keys\" class=\"link-primary\" target=\"_blank\">Or get your own OpenAI API key here !</a>",unsafe_allow_html=True)

    # Get available models (will use fallback if no API key yet)
    model_options = lib.get_available_models(api_key=MY_API_KEY if MY_API_KEY and MY_API_KEY.startswith('sk-') else None)

    # Flatten the options for selectbox
    all_models = []
    for category, models in model_options.items():
        for model in models:
            all_models.append(f"{category}: {model}")

    # Select model with categorization
    selected_model_with_category = st.selectbox(
        "**Select GPT model**",
        all_models,
        index=0,
        help="Choose the OpenAI model you want to use. More capable models provide better results but may cost more."
    )

    # Extract just the model name
    selected_model = selected_model_with_category.split(": ")[-1]
    
    # Check if model is deprecated
    is_deprecated, replacement = lib.check_model_deprecation(selected_model)
    if is_deprecated:
        st.warning(f"⚠️ **{selected_model}** is deprecated. Consider using **{replacement}** instead.", icon="⚠️")
    
    # Show model description
    model_description = lib.get_model_description(selected_model)
    st.caption(model_description)
    
    # Show last update info
    if "_model_cache" in dir(lib) and lib._model_cache["models"] is not None:
        import datetime
        last_update = datetime.datetime.fromtimestamp(lib._model_cache["timestamp"]).strftime("%H:%M:%S")
        st.caption(f"ℹ️ Model list updated at {last_update}")
    else:
        st.caption("ℹ️ Using fallback model list - will update after API key validation")

    # Store in session state
    if "openai_model" not in st.session_state:
        st.session_state["openai_model"] = selected_model
    elif st.session_state["openai_model"] != selected_model:
        st.session_state["openai_model"] = selected_model
    

    if st.button("Submit Your API Key",key="submit_your_api_key_btn1",type="primary"):
        # Basic format validation
        if not MY_API_KEY or not MY_API_KEY.startswith('sk-'):
            st.error("Please enter a valid OpenAI API key (should start with 'sk-')", icon="⚠️")
        else:
            client = OpenAI(api_key=MY_API_KEY)
            try:
                # Simple validation with minimal API call
                with st.spinner("Validating API key and fetching latest models..."):
                    # Just make a simple models list call - much faster than completion
                    models = client.models.list()
                    
                    # Refresh model list with the new API key
                    lib.get_available_models(api_key=MY_API_KEY, force_refresh=True)
                    
                    # Check if selected model is available (without making a completion call)
                    available_models = [model.id for model in models.data]
                    if selected_model not in available_models and not any(model_id.startswith(selected_model) for model_id in available_models):
                        # Try to use the default model
                        default_model = lib.get_default_model()
                        if default_model in available_models:
                            st.warning(f"Model '{selected_model}' may not be available. Using {default_model} instead.", icon="⚠️")
                            selected_model = default_model
                        else:
                            # Fallback to gpt-3.5-turbo
                            st.warning(f"Model '{selected_model}' may not be available. Falling back to gpt-3.5-turbo.", icon="⚠️")
                            selected_model = "gpt-3.5-turbo"
                        st.session_state["openai_model"] = selected_model
                    
                    # Store API key immediately - no need for test completion
                    st.session_state["gpt_api_key"] = MY_API_KEY
                    st.success(f"✅ API key validated! Using model: {st.session_state['openai_model']}", icon="✅")
                    st.info(f"🔄 Model list updated with {len([m for models in lib._model_cache['models'].values() for m in models])} available models")
                    
                    # Load data sources in background (non-blocking)
                    st.rerun()
                    
            except openai.AuthenticationError:
                st.error("🔑 **Invalid API Key** - Please check your OpenAI API key and try again.", icon="⚠️")
                st.info("💡 **Need help?** Get your API key at: https://platform.openai.com/api-keys")
            except openai.PermissionDeniedError:
                st.error("🚫 **Access Denied** - Your API key doesn't have permission to access this model.", icon="⚠️")
                st.info("💡 **Try:** Select a different model or check your OpenAI account permissions.")
            except openai.RateLimitError:
                st.error("⏱️ **Rate Limit Reached** - Too many requests. Please wait a moment and try again.", icon="⚠️")
            except openai.APIConnectionError:
                st.error("🌐 **Connection Error** - Unable to connect to OpenAI. Please check your internet connection.", icon="⚠️")
            except Exception as e:
                # For any other unexpected errors, show a generic friendly message
                st.error("❌ **Something went wrong** - Please check your API key and try again.", icon="⚠️")
                st.info("💡 **Need help?** Make sure your API key starts with 'sk-' and is valid.")
else:
    st.header(f"ChatGPT Discussion ({st.session_state["openai_model"] })")

    # Initialize session state variables
    # Initialize messages in session state if they don't exist
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Initialize processed_anonymizations if it doesn't exist
    if "processed_anonymizations" not in st.session_state:
        st.session_state.processed_anonymizations = set()
    
    # Initialize conversation persistence variables
    if "current_conversation_id" not in st.session_state:
        st.session_state.current_conversation_id = None
    
    if "last_conversation_save_time" not in st.session_state:
        st.session_state.last_conversation_save_time = None
    
    if "last_save_message_count" not in st.session_state:
        st.session_state.last_save_message_count = 0
    
    if "conversation_title" not in st.session_state:
        # Auto-generate title from document if available
        if "saved_anonymisation" in st.session_state:
            st.session_state.conversation_title = f"Chat: {st.session_state['saved_anonymisation']['Title']}"
        else:
            st.session_state.conversation_title = None
    
    # Auto-save configuration
    if "auto_save_enabled" not in st.session_state:
        st.session_state.auto_save_enabled = True
    
    if "auto_save_interval" not in st.session_state:
        st.session_state.auto_save_interval = 2  # Save every 2 messages (1 Q&A exchange)

    # Initialize client
    client = OpenAI(api_key=st.session_state["gpt_api_key"])

    # Load data sources on first access to chat area (lazy loading)
    if "data_sources_loaded" not in st.session_state:
        with st.spinner("Loading your data sources..."):
            # Check for existing documents in the persistent storage
            persist_directory = './chroma_db'
            if os.path.exists(persist_directory) and os.path.isdir(persist_directory):
                document_metadata = lib.load_document_metadata(persist_directory)
                if document_metadata:
                    try:
                        # Initialize ChromaDB client with custom embedding function
                        if hasattr(lib, 'RAG_AVAILABLE') and lib.RAG_AVAILABLE:
                            # Create custom embedding function
                            embedding_function = lib.OpenAIEmbeddingFunction(api_key=st.session_state["gpt_api_key"])
                            
                            # Create ChromaDB client
                            import chromadb
                            from chromadb.config import Settings
                            
                            chroma_client = chromadb.PersistentClient(
                                path=persist_directory,
                                settings=Settings(
                                    anonymized_telemetry=False,
                                    allow_reset=True
                                )
                            )
                            
                            # Connect to the collection
                            collection = chroma_client.get_collection(
                                name="document_collection", 
                                embedding_function=embedding_function
                            )
                            
                            # Store in session state
                            st.session_state.vector_db = collection
                            st.session_state.selected_doc_sources = list(document_metadata.keys())
                            
                            # If there are documents available, select the first one as the active document
                            # This allows users to immediately use their saved documents
                            if document_metadata and "saved_anonymisation" not in st.session_state:
                                # Get the most recent document (sorted by timestamp)
                                doc_list = []
                                for doc_hash, doc_data in document_metadata.items():
                                    doc_list.append({
                                        "hash": doc_hash,
                                        "title": doc_data["title"],
                                        "timestamp": doc_data.get("timestamp", 0)
                                    })
                                # Sort by timestamp (newest first)
                                doc_list.sort(key=lambda x: x["timestamp"], reverse=True)
                                
                                if doc_list:
                                    # Set the latest document as the active document
                                    latest_doc = doc_list[0]
                                    
                                    # Get document hash to retrieve full metadata
                                    doc_hash = latest_doc["hash"]
                                    
                                    # Get entities from metadata if available
                                    entities_data = {
                                        "Text": [],
                                        "Replacement": [],
                                        "Category": []
                                    }
                                    
                                    # Check if this document has stored entities
                                    if 'entities' in document_metadata[doc_hash]:
                                        entities_data = document_metadata[doc_hash]['entities']
                                    
                                    st.session_state["saved_anonymisation"] = {
                                        "Title": latest_doc["title"],
                                        "Time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                        "Entities": entities_data  # Use stored entities if available
                                    }
                                    
                                    # Also initialize saved_gpt_answers with the same document
                                    # This ensures the revert page will have access to the document
                                    st.session_state["saved_gpt_answers"] = {
                                        "Title": latest_doc["title"],
                                        "Time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                        "Data": "",  # Will be populated with conversation content
                                        "Entities": entities_data,  # Use stored entities if available
                                        "Attendees": {}
                                    }
                    except Exception as e:
                        st.error(f"Error loading document database: {e}")
            
            # Load existing tabular datasets
            try:
                persistent_datasets = lib.load_tabular_datasets()
                if persistent_datasets:
                    # Merge with existing session datasets instead of overwriting
                    if "tabular_datasets" not in st.session_state:
                        st.session_state.tabular_datasets = {}
                    
                    # Add debug info if debug mode is enabled
                    if st.session_state.get("show_debug", False):
                        st.write(f"🔧 DEBUG: Loading {len(persistent_datasets)} persistent datasets")
                        st.write(f"🔧 DEBUG: Current session has {len(st.session_state.tabular_datasets)} datasets")
                    
                    # Add persistent datasets, but avoid duplicates
                    for dataset_name, df in persistent_datasets.items():
                        if dataset_name not in st.session_state.tabular_datasets:
                            st.session_state.tabular_datasets[dataset_name] = df
                            if st.session_state.get("show_debug", False):
                                st.write(f"🔧 DEBUG: Added persistent dataset: {dataset_name}")
                        elif st.session_state.get("show_debug", False):
                            st.write(f"🔧 DEBUG: Skipped duplicate dataset: {dataset_name}")
            except Exception as e:
                st.error(f"Error loading tabular datasets: {e}")
            
            # Mark data sources as loaded
            st.session_state.data_sources_loaded = True



    # Auto-save conversation function
    def auto_save_conversation():
        """
        Auto-save the current conversation if conditions are met.
        Saves every N messages based on auto_save_interval setting.
        """
        if not st.session_state.get("auto_save_enabled", True):
            return
        
        messages = st.session_state.get("messages", [])
        if len(messages) == 0:
            return
        
        # Check if we should save based on message count
        last_save_count = st.session_state.get("last_save_message_count", 0)
        auto_save_interval = st.session_state.get("auto_save_interval", 2)
        
        # Auto-save if we've added enough new messages
        if len(messages) - last_save_count >= auto_save_interval:
            try:
                conversation_id = lib.save_conversation(
                    messages=messages,
                    conversation_id=st.session_state.get("current_conversation_id"),
                    title=st.session_state.get("conversation_title"),
                    auto_save=True
                )
                
                # Update session state
                st.session_state.current_conversation_id = conversation_id
                st.session_state.last_conversation_save_time = time.time()
                st.session_state.last_save_message_count = len(messages)
                
                # Show subtle toast notification for auto-save
                st.toast("💾 Conversation auto-saved", icon="✅")
                
                if st.session_state.get("show_debug", False):
                    st.success(f"💾 Auto-saved conversation ({len(messages)} messages)")
                    
            except Exception as e:
                if st.session_state.get("show_debug", False):
                    st.error(f"Auto-save failed: {e}")
    
    # Manual save conversation function
    def manual_save_conversation():
        """Manually save the current conversation"""
        messages = st.session_state.get("messages", [])
        if len(messages) == 0:
            st.warning("No messages to save")
            return False
        
        try:
            conversation_id = lib.save_conversation(
                messages=messages,
                conversation_id=st.session_state.get("current_conversation_id"),
                title=st.session_state.get("conversation_title"),
                auto_save=False
            )
            
            # Update session state
            st.session_state.current_conversation_id = conversation_id
            st.session_state.last_conversation_save_time = time.time()
            st.session_state.last_save_message_count = len(messages)
            
            st.success(f"💾 Conversation saved successfully! ({len(messages)} messages)")
            return True
            
        except Exception as e:
            st.error(f"Failed to save conversation: {e}")
            return False

    # Conversation Browser Dialog Function
    @st.dialog("💬 Saved Conversations", width="large")
    def show_conversation_browser():
        st.markdown("**Browse and restore your saved conversations**")
        
        # Get list of saved conversations
        conversations = lib.list_conversations(limit=50)
        
        if not conversations:
            st.info("📭 No saved conversations found. Conversations are auto-saved as you chat.")
            st.markdown("💡 **Tip:** Conversations are automatically saved every 5 messages, or you can manually save using the **💾 Save** button.")
            return
        
        st.success(f"Found {len(conversations)} saved conversation(s)")
        
        # Display conversations
        for conv in conversations:
            with st.expander(f"💬 {conv['title']}", expanded=False):
                col1, col2, col3 = st.columns([2, 1, 1])
                
                with col1:
                    st.write(f"**Messages:** {conv['message_count']}")
                    st.write(f"**Model:** {conv['model']}")
                
                with col2:
                    # Parse the date
                    try:
                        from datetime import datetime
                        saved_date = datetime.fromisoformat(conv['last_saved'])
                        st.write(f"**Saved:** {saved_date.strftime('%Y-%m-%d')}")
                        st.write(f"**Time:** {saved_date.strftime('%H:%M')}")
                    except:
                        st.write(f"**Saved:** {conv.get('last_saved', 'Unknown')[:10]}")
                
                with col3:
                    # Action buttons row 1
                    col_load, col_fork, col_rename = st.columns(3)
                    
                    with col_load:
                        # Load button
                        if st.button("📂", key=f"load_{conv['id']}", use_container_width=True, help="Load conversation"):
                            # Check for unsaved changes
                            save_status = lib.get_conversation_save_status()
                            if save_status["has_unsaved"] and len(st.session_state.messages) > 0:
                                st.warning("⚠️ You have unsaved changes in your current conversation!")
                                if st.button("⚠️ Load anyway (lose unsaved)", key=f"confirm_load_{conv['id']}", type="secondary"):
                                    load_conversation_data(conv['id'])
                            else:
                                load_conversation_data(conv['id'])
                    
                    with col_fork:
                        # Fork button (create numbered copy)
                        if st.button("🔀", key=f"fork_{conv['id']}", use_container_width=True, help="Create numbered copy"):
                            fork_conversation(conv['id'], conv['title'])
                    
                    with col_rename:
                        # Rename button
                        if st.button("✏️", key=f"rename_btn_{conv['id']}", use_container_width=True, help="Rename conversation"):
                            st.session_state[f"renaming_{conv['id']}"] = True
                            st.rerun()
                    
                    # Rename input (shown when rename button clicked)
                    if st.session_state.get(f"renaming_{conv['id']}", False):
                        st.markdown("**✏️ Rename Conversation:**")
                        new_title = st.text_input(
                            "New title:",
                            value=conv['title'],
                            key=f"rename_input_{conv['id']}",
                            label_visibility="collapsed"
                        )
                        
                        col_save, col_cancel = st.columns(2)
                        with col_save:
                            if st.button("💾 Save", key=f"rename_save_{conv['id']}", use_container_width=True):
                                if new_title and new_title.strip():
                                    # Check if title already exists (excluding current conversation)
                                    unique_title = lib.get_unique_conversation_title(new_title, exclude_id=conv['id'])
                                    
                                    if lib.rename_conversation(conv['id'], unique_title):
                                        st.success(f"✅ Renamed to: {unique_title}")
                                        st.session_state[f"renaming_{conv['id']}"] = False
                                        st.rerun()
                                    else:
                                        st.error("Failed to rename")
                                else:
                                    st.error("Title cannot be empty")
                        
                        with col_cancel:
                            if st.button("❌ Cancel", key=f"rename_cancel_{conv['id']}", use_container_width=True):
                                st.session_state[f"renaming_{conv['id']}"] = False
                                st.rerun()
                    
                    # Delete button with confirmation
                    pending_delete_key = f"pending_delete_{conv['id']}"
                    
                    if st.session_state.get(pending_delete_key, False):
                        # Show confirmation
                        st.error(f"⚠️ Confirm delete?")
                        col_a, col_b = st.columns(2)
                        with col_a:
                            if st.button("✅ Yes, Delete", key=f"confirm_yes_{conv['id']}", type="secondary", use_container_width=True):
                                if lib.delete_conversation(conv['id']):
                                    st.success(f"Deleted: {conv['title']}")
                                    st.session_state[pending_delete_key] = False
                                    st.rerun()
                                else:
                                    st.error("Failed to delete")
                                    st.session_state[pending_delete_key] = False
                        with col_b:
                            if st.button("❌ Cancel", key=f"confirm_no_{conv['id']}", use_container_width=True):
                                st.session_state[pending_delete_key] = False
                                st.rerun()
                    else:
                        # Show delete button
                        if st.button("🗑️ Delete", key=f"delete_{conv['id']}", use_container_width=True):
                            st.session_state[pending_delete_key] = True
                            st.rerun()
    
    def fork_conversation(conversation_id, original_title):
        """Create a numbered copy of an existing conversation"""
        conv_data = lib.load_conversation(conversation_id)
        
        if conv_data:
            # Generate unique numbered title
            new_title = lib.get_unique_conversation_title(original_title)
            
            # Save as new conversation with numbered title
            new_conv_id = lib.save_conversation(
                messages=conv_data.get("messages", []),
                conversation_id=None,  # Force new ID
                title=new_title,
                auto_save=False
            )
            
            st.success(f"✅ Created: {new_title}")
            st.rerun()
        else:
            st.error("Failed to fork conversation")
    
    def load_conversation_data(conversation_id):
        """Load a conversation and restore it to the current session"""
        conv_data = lib.load_conversation(conversation_id)
        
        if conv_data:
            # Restore messages
            loaded_messages = conv_data.get("messages", [])
            original_count = len(loaded_messages)
            
            # Apply context window management if conversation is large
            # This prevents token overflow when loading old conversations
            if original_count > 30:
                # Manage context to prevent token overflow
                loaded_messages = lib.manage_context_window(
                    loaded_messages,
                    max_context_messages=30,  # Keep last 30 messages
                    preserve_recent=20  # Always preserve last 20
                )
                
                # Show info about context management
                trimmed_count = original_count - len(loaded_messages)
                if trimmed_count > 0:
                    st.info(f"ℹ️ Loaded conversation with {original_count} messages. {trimmed_count} older messages were summarized to optimize performance.")
            
            st.session_state.messages = loaded_messages
            st.session_state.original_message_count = original_count  # Track original size
            
            # Restore conversation metadata
            st.session_state.current_conversation_id = conv_data.get("id")
            st.session_state.conversation_title = conv_data.get("title")
            st.session_state.last_conversation_save_time = time.time()
            st.session_state.last_save_message_count = len(st.session_state.messages)
            
            # Restore charts if any
            if "stored_charts" in conv_data:
                st.session_state.stored_charts = conv_data["stored_charts"]
            
            st.success(f"✅ Loaded conversation: {conv_data.get('title')}")
            st.rerun()
        else:
            st.error("Failed to load conversation")

    # Model Settings Dialog Function
    @st.dialog("⚙️ Model Settings", width="large")
    def show_model_settings():
        st.markdown("**Configure AI model settings and advanced options**")
        
        # AI Model Configuration Section with Expander
        with st.expander("🤖 AI Model Configuration", expanded=True):
            # Display current model and provide model information
            current_model = st.session_state['openai_model']
            st.info(f"Current model: **{current_model}**")
            
            # Check if current model is deprecated
            is_deprecated, replacement = lib.check_model_deprecation(current_model)
            if is_deprecated:
                st.warning(f"⚠️ This model is deprecated. Recommended replacement: **{replacement}**")
            
            # Show model description
            model_description = lib.get_model_description(current_model)
            st.caption(model_description)
            
            # Model rate information
            st.caption("**Note:** Different models have different pricing. Check [OpenAI API Pricing](https://openai.com/api/pricing/) for details.")
            
            # Show model list status
            col1, col2 = st.columns(2)
            with col1:
                # Add a button to reset API key and change model
                if st.button("Change AI Model", use_container_width=True):
                    del st.session_state["gpt_api_key"]
                    del st.session_state["openai_model"]
                    # Clear vector DB session state if it exists
                    if "vector_db" in st.session_state:
                        del st.session_state["vector_db"]
                    st.rerun()
            
            with col2:
                # Add a button to refresh model list
                if st.button("🔄 Refresh Models", use_container_width=True, help="Fetch latest models from OpenAI API"):
                    if "gpt_api_key" in st.session_state:
                        lib.get_available_models(api_key=st.session_state["gpt_api_key"], force_refresh=True)
                        st.success("✅ Model list refreshed!")
                        st.rerun()
            
            # Show cache status
            if lib._model_cache["models"] is not None:
                import datetime
                last_update = datetime.datetime.fromtimestamp(lib._model_cache["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
                st.caption(f"📊 Model list last updated: {last_update}")
            else:
                st.caption("📊 Using fallback model list")
        
        # Advanced Settings Section with Expander
        with st.expander("🔧 Advanced Settings", expanded=False):
            # Debug mode toggle
            debug_enabled = st.checkbox(
                "Enable debug mode", 
                key="debug_toggle", 
                value=st.session_state.get("show_debug", False),
                help="Show detailed information about agent tool usage and execution steps. Useful for understanding how the AI processes your questions."
            )
            st.session_state.show_debug = debug_enabled
            
            # Max iterations configuration
            max_iterations = st.slider(
                "Max Analysis Steps",
                min_value=3,
                max_value=20,
                value=st.session_state.get("max_iterations", 10),
                help="Maximum number of tool calls the AI agent can make per response. Higher values allow more complex analysis but may take longer."
            )
            st.session_state.max_iterations = max_iterations
            
            if debug_enabled:
                st.info("🔍 Debug mode is **ON**. You'll see detailed execution information after each AI response.")
            else:
                st.info("🔍 Debug mode is **OFF**. Clean interface with no technical details shown.")
            
            # Context Window Management Settings
            st.divider()
            st.markdown("**Context Window Management**")
            
            # Max context messages setting
            max_context_messages = st.slider(
                "Max Context Messages",
                min_value=10,
                max_value=50,
                value=st.session_state.get("max_context_messages", 20),
                help="Maximum number of messages to keep in full before summarizing older ones"
            )
            st.session_state.max_context_messages = max_context_messages
            
            # Preserve recent messages setting
            preserve_recent = st.slider(
                "Recent Messages to Preserve",
                min_value=5,
                max_value=15,
                value=st.session_state.get("preserve_recent", 10),
                help="Number of most recent messages to always keep in full context"
            )
            st.session_state.preserve_recent = preserve_recent
            
            # Tool output truncation setting
            tool_output_max = st.slider(
                "Tool Output Max Length",
                min_value=500,
                max_value=2000,
                value=st.session_state.get("tool_output_max", 1000),
                help="Maximum length for tool outputs before truncation (saves context space)"
            )
            st.session_state.tool_output_max = tool_output_max
            
            # Show current configuration
            st.caption(f"Current settings: Max steps = {max_iterations}, Debug = {'ON' if debug_enabled else 'OFF'}")
            st.caption(f"Context settings: {max_context_messages} max messages, {preserve_recent} recent preserved, {tool_output_max} tool output limit")
        
        # System Information Section with Expander
        with st.expander("ℹ️ System Information", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Session Documents", len(st.session_state.get('selected_doc_sources', [])))
                st.metric("Session Datasets", len(st.session_state.get('tabular_datasets', {})))
            
            with col2:
                st.metric("Chat Messages", len(st.session_state.get('messages', [])))
                st.metric("Stored Charts", len(st.session_state.get('stored_charts', {})))
        
        # Session Management Section with Expander
        with st.expander("🗃️ Session Management", expanded=False):
            col1, col2 = st.columns(2)
            with col1:
                if st.button("Clear Chat History", use_container_width=True):
                    st.session_state.messages = []
                    if "stored_charts" in st.session_state:
                        st.session_state.stored_charts = {}
                    if "current_chart_id" in st.session_state:
                        st.session_state.current_chart_id = None
                    if "current_chart_ids" in st.session_state:
                        st.session_state.current_chart_ids = []
                    st.success("Chat history cleared!")
            
            with col2:
                if st.button("Reset All Settings", use_container_width=True):
                    # Reset to defaults
                    st.session_state.show_debug = False
                    st.session_state.max_iterations = 10
                    st.session_state.max_context_messages = 20
                    st.session_state.preserve_recent = 10
                    st.session_state.tool_output_max = 1000
                    st.success("All settings reset to defaults!")

    # Process document and create vector database if not already processed
    if "saved_anonymisation" in st.session_state:
        # Generate a unique key for this anonymization
        anon_key = st.session_state["saved_anonymisation"]["Title"] + "_" + st.session_state["saved_anonymisation"]["Time"]
        
        # Check if this anonymization has already been processed
        if anon_key not in st.session_state.processed_anonymizations and "vector_db" in st.session_state:
            # We have a new document to add to the existing vector DB
            # Only process if we have actual document data
            if "Data" in st.session_state["saved_anonymisation"] and st.session_state["saved_anonymisation"]["Data"]:
                with st.spinner(f"Adding new document '{st.session_state['saved_anonymisation']['Title']}' to the database..."):
                    # Get document info
                    document_title = st.session_state["saved_anonymisation"]["Title"]
                    document_content = st.session_state["saved_anonymisation"]["Data"]
                    document_entities = st.session_state["saved_anonymisation"]["Entities"]
                    
                    # Add to vector database
                    st.session_state.vector_db = lib.create_vector_db_from_text(
                        document_content, 
                        document_title, 
                        st.session_state["gpt_api_key"],
                        document_entities  # Pass entities for reverse anonymization
                    )
                    
                    if st.session_state.vector_db is not None:
                        st.success(f"Document '{document_title}' added successfully to the database!")
                        # Mark as processed
                        st.session_state.processed_anonymizations.add(anon_key)
                        
                        # Update selected_doc_sources to reflect the new document
                        persist_directory = './chroma_db'
                        document_metadata = lib.load_document_metadata(persist_directory)
                        if document_metadata:
                            st.session_state.selected_doc_sources = list(document_metadata.keys())
                    else:
                        st.error(f"Failed to add document '{document_title}' to the database.")
                
        elif anon_key not in st.session_state.processed_anonymizations and "vector_db" not in st.session_state:
            # This is the first document, and we need to create the vector DB
            # Only process if we have actual document data
            if "Data" in st.session_state["saved_anonymisation"] and st.session_state["saved_anonymisation"]["Data"]:
                with st.spinner("Processing your document... This might take a minute."):
                    # Check if RAG is available
                    if not hasattr(lib, 'RAG_AVAILABLE') or not lib.RAG_AVAILABLE:
                        st.warning("RAG functionality is not available. The app will use the traditional chunking approach instead.")
                        # Set vector_db to None to skip RAG-related code
                        st.session_state.vector_db = None
                    else:
                        # Create vector database from anonymized content
                        document_title = st.session_state["saved_anonymisation"]["Title"]
                        document_content = st.session_state["saved_anonymisation"]["Data"]
                        document_entities = st.session_state["saved_anonymisation"]["Entities"]
                        
                        # Create vector database
                        st.session_state.vector_db = lib.create_vector_db_from_text(
                            document_content, 
                            document_title, 
                            st.session_state["gpt_api_key"],
                            document_entities  # Pass entities for reverse anonymization
                        )
                        
                        if st.session_state.vector_db is not None:
                            st.success(f"Document processed successfully! You can now ask questions about '{document_title}'")
                            # Mark as processed
                            st.session_state.processed_anonymizations.add(anon_key)
                            
                            # Update selected_doc_sources to reflect the new document
                            persist_directory = './chroma_db'
                            document_metadata = lib.load_document_metadata(persist_directory)
                            if document_metadata:
                                st.session_state.selected_doc_sources = list(document_metadata.keys())
                        else:
                            st.error("Failed to process document with RAG. Falling back to traditional approach.")
                            # Set vector_db to None to indicate RAG is not available
                            st.session_state.vector_db = None

    # Show unsaved changes warning banner (prominent at top)
    if len(st.session_state.messages) > 0:
        save_status = lib.get_conversation_save_status()
        if save_status["has_unsaved"]:
            unsaved_count = save_status["message_count"] - save_status["last_save_message_count"]
            st.warning(f"⚠️ **Unsaved Changes:** You have {unsaved_count} message(s) not saved yet. Use the **💾 Save** button below to save your conversation.", icon="⚠️")
            
            # Inject JavaScript warning for browser navigation
            lib.inject_navigation_warning(has_unsaved_changes=True, unsaved_count=unsaved_count)
    
    # Check available data sources and show appropriate messages
    has_documents = ("vector_db" in st.session_state and st.session_state.vector_db is not None) or ("saved_anonymisation" in st.session_state)
    has_datasets = "tabular_datasets" in st.session_state and st.session_state.tabular_datasets
    
    if not has_documents and not has_datasets:
        st.warning("No data sources loaded. Please upload documents or datasets to begin analysis.")
        st.page_link(page="pages/00_home.py", label="Go To Home Page", icon=":material/home:", use_container_width=True)
    else:
        # Show what's available in a compact info box with active/inactive status
        available_sources = []
        if has_documents:
            # Get document counts from metadata
            persist_directory = './chroma_db'
            document_metadata = lib.load_document_metadata(persist_directory)
            if document_metadata:
                active_docs = sum(1 for doc_data in document_metadata.values() if doc_data.get("active", True))
                total_docs = len(document_metadata)
                
                if active_docs == total_docs:
                    available_sources.append(f"{total_docs} text document(s)")
                else:
                    available_sources.append(f"{active_docs}/{total_docs} text document(s)")
            elif "saved_anonymisation" in st.session_state:
                # If we have a current document but no persistent storage yet, count it as 1
                available_sources.append("1 text document")
        
        if has_datasets:
            # Get dataset counts from metadata
            tabular_metadata = lib.get_tabular_metadata()
            if tabular_metadata:
                active_datasets = sum(1 for dataset_data in tabular_metadata.values() if dataset_data.get("active", True))
                total_datasets = len(tabular_metadata)
                
                if active_datasets == total_datasets:
                    available_sources.append(f"{total_datasets} tabular dataset(s)")
                else:
                    available_sources.append(f"{active_datasets}/{total_datasets} tabular dataset(s)")
        
        # Show data sources, context optimization, and save status
        col1, col2, col3 = st.columns([2, 1, 1])
        with col1:
            if available_sources:
                sources_text = ', '.join(available_sources)
                st.info(f"📊 Ready to analyze: {sources_text}")
            else:
                st.info("📊 Ready to analyze your data sources")
        
        with col2:
            # Show context optimization status
            if len(st.session_state.messages) > st.session_state.get("max_context_messages", 20):
                context_stats = lib.get_context_stats(st.session_state.messages)
                st.info(f"🧠 Context: {context_stats['estimated_tokens']:,} tokens")
            elif len(st.session_state.messages) > 10:
                context_stats = lib.get_context_stats(st.session_state.messages)
                st.success(f"🧠 Context: {context_stats['estimated_tokens']:,} tokens")
        
        with col3:
            # Show conversation save status
            save_status = lib.get_conversation_save_status()
            if save_status["last_save_time"]:
                time_since_save = int((time.time() - save_status["last_save_time"]) / 60)  # Minutes
                if save_status["has_unsaved"]:
                    unsaved_count = save_status["message_count"] - save_status["last_save_message_count"]
                    st.warning(f"💾 {unsaved_count} unsaved")
                else:
                    if time_since_save == 0:
                        st.success("💾 Just saved")
                    else:
                        st.success(f"💾 Saved {time_since_save}m ago")
            elif len(st.session_state.messages) > 0:
                st.warning(f"💾 Not saved")

    # Welcome message for first-time users (when no chat history exists)
    if not st.session_state.messages:
        # Check if user explicitly chose to start new chat
        if st.session_state.get("skip_conversation_welcome", False):
            # User clicked "Start New Chat" - skip the welcome screen
            st.session_state.skip_conversation_welcome = False
            # Show nothing, let user start typing
            pass
        else:
            # Check if there are saved conversations
            saved_conversations = lib.list_conversations(limit=3)
            
            if saved_conversations:
                # Show resume/start new interface when saved conversations exist
                with st.chat_message("assistant"):
                    st.markdown("### 💬 Welcome Back!")
                    st.markdown(f"You have **{len(saved_conversations)}** saved conversation(s). Would you like to resume or start fresh?")
                    
                    # Action buttons
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        if st.button("📂 Resume Last Session", use_container_width=True, type="primary"):
                            # Load the most recent conversation
                            most_recent = saved_conversations[0]
                            load_conversation_data(most_recent['id'])
                    
                    with col2:
                        if st.button("✨ Start New Chat", use_container_width=True):
                            # Clear any loaded conversation state and start fresh
                            st.session_state.skip_conversation_welcome = True
                            # Clear conversation tracking to start completely fresh
                            st.session_state.current_conversation_id = None
                            st.session_state.conversation_title = None
                            st.session_state.last_save_message_count = 0
                            st.rerun()
                    
                    st.markdown("---")
                    st.markdown("**📚 Recent Conversations:**")
                    
                    # Show list of recent conversations
                    for idx, conv in enumerate(saved_conversations, 1):
                        with st.container():
                            col1, col2, col3 = st.columns([3, 2, 1])
                            
                            with col1:
                                st.markdown(f"**{idx}. {conv['title']}**")
                            
                            with col2:
                                # Format timestamp
                                try:
                                    from datetime import datetime
                                    saved_date = datetime.fromisoformat(conv['last_saved'])
                                    time_ago = datetime.now() - saved_date
                                    
                                    if time_ago.days > 0:
                                        time_str = f"{time_ago.days}d ago"
                                    elif time_ago.seconds >= 3600:
                                        time_str = f"{time_ago.seconds // 3600}h ago"
                                    elif time_ago.seconds >= 60:
                                        time_str = f"{time_ago.seconds // 60}m ago"
                                    else:
                                        time_str = "Just now"
                                    
                                    st.caption(f"{conv['message_count']} msgs • {time_str}")
                                except:
                                    st.caption(f"{conv['message_count']} messages")
                            
                            with col3:
                                if st.button("📂", key=f"load_welcome_{conv['id']}", help="Load this conversation"):
                                    load_conversation_data(conv['id'])
                    
                    st.markdown("---")
                    st.markdown("💡 **Tip:** Click **💬 Conversations** button below to see all saved conversations")
            else:
                # Show normal welcome message when no saved conversations exist
                with st.chat_message("assistant"):
                    st.markdown("### Hello! How can I help you today?")
                    
                    # Create a concise welcome message based on available data sources
                    if has_documents and has_datasets:
                        welcome_msg = """I can analyze your **documents** and **datasets** to help you:
• Summarize key insights • Create charts and visualizations • Extract action items • Generate reports

💡 **Try asking:** *"What are the key insights from my data?"* or *"Create a summary with charts"* or click the **Prompt Ideas** button below"""
                    
                    elif has_documents:
                        welcome_msg = """I can analyze your **text documents** to help you:
• Summarize key points • Extract action items • Answer specific questions • Create reports

💡 **Try asking:** *"What were the main decisions?"* or *"Create a summary of key themes"* or click the **Prompt Ideas** button below"""
                    
                    elif has_datasets:
                        welcome_msg = """I can analyze your **tabular data** to help you:
• Create charts and visualizations • Analyze trends and patterns • Generate insights

💡 **Try asking:** *"Show me the top 10 items"* or *"Create a trend chart over time"* or click the **Prompt Ideas** button below"""
                    
                    else:
                        welcome_msg = """I can help you analyze documents and datasets! 
🚀 **To get started:** Upload data on the Home page, then ask me questions about it or click the **Prompt Ideas** button below"""
                    
                    st.markdown(welcome_msg)

    # ===== CONVERSATION TOKEN MANAGEMENT & WARNINGS =====
    # Check if we need to warn about token limits or apply automatic trimming
    current_model = st.session_state.get("openai_model", "gpt-3.5-turbo")
    context_stats = lib.get_context_stats(st.session_state.messages)
    
    # Get model-specific token limits
    limit_info = lib.get_model_token_limit(current_model)
    estimated_tokens = context_stats["estimated_tokens"]
    
    # Show warning if approaching token limits
    if estimated_tokens > limit_info["warning_threshold"]:
        warning_pct = int((estimated_tokens / limit_info["limit"]) * 100)
        st.warning(f"⚠️ **Token Limit Warning**: Your conversation is using ~{estimated_tokens:,} tokens ({warning_pct}% of {limit_info['name']}'s {limit_info['limit']:,} token limit). Consider starting a new conversation or the AI may not have full context.")
    
    # Auto-trim if conversation is getting very large (90% of limit)
    auto_trim_threshold = int(limit_info["limit"] * 0.9)
    if estimated_tokens > auto_trim_threshold and len(st.session_state.messages) > 20:
        # Apply automatic context management
        original_count = len(st.session_state.messages)
        st.session_state.messages = lib.manage_context_window(
            st.session_state.messages,
            max_context_messages=20,
            preserve_recent=15
        )
        trimmed_count = original_count - len(st.session_state.messages)
        if trimmed_count > 0:
            st.info(f"🔄 **Auto-optimization**: Conversation trimmed to prevent token overflow. {trimmed_count} older messages summarized.")
    
    # ===== UI DISPLAY TRUNCATION =====
    # For better performance, only display recent messages if conversation is very long
    messages_to_display = st.session_state.messages
    display_limit = 50  # Show last 50 messages by default
    show_all_messages = st.session_state.get("show_all_messages", False)
    
    if len(st.session_state.messages) > display_limit and not show_all_messages:
        # Show info about hidden messages
        hidden_count = len(st.session_state.messages) - display_limit
        st.info(f"📜 Showing last {display_limit} messages ({hidden_count} older messages hidden for performance)")
        
        col1, col2 = st.columns([1, 4])
        with col1:
            if st.button("📖 Show Full History", key="show_full_history"):
                st.session_state.show_all_messages = True
                st.rerun()
        with col2:
            st.caption(f"Total conversation: {len(st.session_state.messages)} messages (~{estimated_tokens:,} tokens)")
        
        # Only display recent messages
        messages_to_display = st.session_state.messages[-display_limit:]
    elif show_all_messages and len(st.session_state.messages) > display_limit:
        # User chose to see all messages, show option to collapse
        st.success(f"📖 Showing all {len(st.session_state.messages)} messages")
        if st.button("📋 Show Recent Only", key="show_recent_only"):
            st.session_state.show_all_messages = False
            st.rerun()
    
    # Display chat messages directly (following Streamlit best practices)
    for message in messages_to_display:
        with st.chat_message(message["role"]):
            # Display the message content
            message_content = message["content"]
            
            # Check if this message contains chart IDs (support multiple charts)
            import re
            chart_matches = re.findall(r'\[CHART_ID:(chart_\d+(?:_\d+)?)\]', message_content)
            
            if st.session_state.get("show_debug", False):
                st.write(f"🔧 DEBUG: Checking message for chart IDs: {message_content[:100]}...")
                if chart_matches:
                    st.write(f"🔧 DEBUG: Found chart IDs in message: {chart_matches}")
                else:
                    st.write(f"🔧 DEBUG: No chart IDs found in message")
            
            if chart_matches and message["role"] == "assistant":
                # Remove all chart IDs from the displayed message
                display_content = re.sub(r'\[CHART_ID:chart_\d+(?:_\d+)?\]', '', message_content)
                st.markdown(display_content)
                
                # Display all charts that exist in stored charts
                for chart_id in chart_matches:
                    if st.session_state.get("show_debug", False):
                        st.write(f"🔧 DEBUG: Looking for chart ID: {chart_id}")
                        st.write(f"🔧 DEBUG: Available chart IDs: {list(st.session_state.get('stored_charts', {}).keys())}")
                    
                    if chart_id in st.session_state.get("stored_charts", {}):
                        chart_config = st.session_state.stored_charts[chart_id]
                        if st.session_state.get("show_debug", False):
                            st.write(f"🔧 DEBUG: Found chart config, recreating chart")
                        
                        fig = lib.recreate_chart_from_config(chart_config)
                        if fig is not None:
                            if st.session_state.get("show_debug", False):
                                st.write(f"🔧 DEBUG: Chart recreation successful, displaying chart")
                            
                            # Anti-aliasing fix: Use specific config to reduce rendering artifacts
                            config = {
                                'displayModeBar': True,
                                'displaylogo': False,
                                'modeBarButtonsToRemove': ['pan2d', 'lasso2d', 'select2d'],
                                'toImageButtonOptions': {
                                    'format': 'svg',  # SVG format for crisp rendering
                                    'filename': 'chart',
                                    'height': 500,
                                    'width': 700,
                                    'scale': 1  # No scaling to prevent artifacts
                                }
                            }
                            st.plotly_chart(fig, use_container_width=True, key=f"current_chart_{chart_id}", config=config)
                            
                            # Collect chart IDs for adding to response
                            chart_ids_for_response = f"\n\n[CHART_ID:{chart_id}]"
                        else:
                            # Show fallback message if chart recreation failed
                            st.warning(f"⚠️ Chart {chart_id} could not be displayed. The chart data may be corrupted or incompatible with the current session.")
                            if st.session_state.get("show_debug", False):
                                st.write("**Debug Info:** Chart config:", chart_config)
                    else:
                        if st.session_state.get("show_debug", False):
                            st.write(f"🔧 DEBUG: Chart ID {chart_id} not found in stored charts")
            else:
                # Display normal message
                st.markdown(message_content)

    # Prompt Assistant Dialog Function
    @st.dialog("💡 Prompt Ideas & Templates", width="large")
    def show_prompt_assistant():
        st.markdown("**Quick access to proven prompt templates for better AI results**")
        
        # Check available data sources for smart suggestions
        has_documents = ("vector_db" in st.session_state and st.session_state.vector_db is not None) or ("saved_anonymisation" in st.session_state)
        has_datasets = "tabular_datasets" in st.session_state and st.session_state.tabular_datasets
        
        # Generate smart templates based on content analysis (zero tokens)
        smart_templates = {}
        if has_documents or has_datasets:
            document_insights = None
            dataset_insights = {}
            
            # Analyze documents locally (zero tokens)
            if "saved_anonymisation" in st.session_state:
                document_content = st.session_state["saved_anonymisation"].get("Data", "")
                document_entities = st.session_state["saved_anonymisation"].get("Entities", {})
                if document_content:
                    document_insights = lib.analyze_document_locally(document_content, document_entities)
            
            # Analyze datasets locally (zero tokens)
            if has_datasets:
                for dataset_name, df in st.session_state.tabular_datasets.items():
                    dataset_insights[dataset_name] = lib.analyze_dataset_locally(df, dataset_name)
            
            # Generate smart templates
            smart_templates = lib.generate_smart_templates_locally(document_insights, dataset_insights)
        
        # Smart suggestions info
        if smart_templates:
            st.success(f"✨ Found **{len(smart_templates)}** personalized suggestions based on your content!")
        elif has_documents and has_datasets:
            st.info("📄📊 Both document and dataset sources are loaded - use any template below!")
        elif has_documents:
            st.info("📄 Document sources loaded - text analysis templates are most relevant")
        elif has_datasets:
            st.info("📊 Dataset sources loaded - data analysis templates are most relevant")
        else:
            st.info("💡 Upload documents or datasets on the Home page to get personalized suggestions")
        
        # Smart Suggestions Section (if available)
        if smart_templates:
            with st.expander("✨ Personalized Suggestions", expanded=False):
                st.markdown("*AI-generated templates based on your actual content*")
                
                for title, template in smart_templates.items():
                    st.markdown(f"**{title}**")
                    with st.container(border=True):
                        st.text_area("Template content:", value=template, height=120, key=f"smart_template_{title}", label_visibility="collapsed")
        
        # Text Analysis Templates Section with Expander
        with st.expander("📊 Text Analysis Templates", expanded=False):
            st.markdown("*Perfect for analyzing documents, transcripts, and text content*")
            
            text_templates = {
                "📝 Meeting Report": """Create a professional meeting report with this structure:

**Executive Summary:** Key outcomes and decisions
**Discussion Points:** Main topics covered  
**Decisions Made:** What was agreed upon
**Action Items:** Who does what by when
**Next Steps:** Follow-up actions

Write in formal, professional tone suitable for organizational reporting.""",
                
                "⚡ Quick Summary": """Provide a concise 3-paragraph summary:

**Paragraph 1:** Main focus and key topics discussed
**Paragraph 2:** Primary conclusions, decisions, or findings  
**Paragraph 3:** Next steps, recommendations, or actions required

Keep each paragraph to 3-4 sentences maximum.""",
                
                "🎯 Action Items": """Extract all actionable elements:

**Immediate Actions:** Tasks for next 1-2 weeks
**Medium-term Actions:** Tasks for next 1-3 months
**Long-term Commitments:** Strategic actions beyond 3 months
**Responsible Parties:** Who is accountable for each action
**Dependencies:** Prerequisites for each action

Format as clear action plan with timelines.""",
                
                "📋 Key Themes": """Identify and organize the main themes:

1. **Theme Identification:** List 3-5 major themes
2. **Content Organization:** Group relevant points under each theme
3. **Key Insights:** Provide 2-3 insights for each theme
4. **Connections:** Note relationships between themes

Present findings in clear, structured format."""
            }
            
            for title, template in text_templates.items():
                st.markdown(f"**{title}**")
                with st.container(border=True):
                    st.text_area("Template content:", value=template, height=150, key=f"text_template_{title}", label_visibility="collapsed")
        
        # Data Analysis Templates Section with Expander
        with st.expander("📈 Data Analysis Templates", expanded=False):
            st.markdown("*Perfect for analyzing CSV files, Excel data, and tabular information*")
            
            data_templates = {
                "📊 Data Overview": """Analyze this dataset and provide:

**Data Structure:** Number of rows, columns, data types
**Key Statistics:** Summary statistics for numerical columns
**Data Quality:** Missing values, outliers, any issues
**Top Insights:** 5 most interesting patterns or findings
**Recommended Charts:** Best visualizations for this data

Be specific about column names and values.""",
                
                "📈 Create Charts": """Create these visualizations from the data:

1. Bar chart showing top 10 values by the main metric
2. Line chart showing trends over time (if time data exists)
3. Summary statistics table for key numerical columns
4. Distribution chart for the most important variable

Include clear titles and labels. Suggest additional charts if helpful.""",
                
                "🔍 Trend Analysis": """Analyze trends and patterns in this data:

**Overall Trends:** General direction and changes over time
**Key Patterns:** Notable increases, decreases, or cyclical patterns  
**Top Performers:** Highest and lowest values with context
**Relationships:** Correlations between different variables
**Time-based Charts:** Visualizations showing temporal patterns

Focus on actionable insights.""",
                
                "⚖️ Compare Groups": """Compare different groups or categories in this data:

**Performance Comparison:** Which groups perform better in key metrics
**Key Differences:** How groups differ across important measures
**Rankings:** Rank groups from best to worst performance
**Visualization:** Create charts highlighting these comparisons
**Insights:** What do these differences tell us

Provide clear, actionable recommendations."""
            }
            
            for title, template in data_templates.items():
                st.markdown(f"**{title}**")
                with st.container(border=True):
                    st.text_area("Template content:", value=template, height=150, key=f"data_template_{title}", label_visibility="collapsed")
        
        # Quick Starters Section with Expander
        with st.expander("🚀 Quick Starters", expanded=False):
            st.markdown("*Simple, ready-to-use prompts for quick analysis*")
            
            quick_templates = {
                "📋 What's in this data?": "Analyze this dataset and tell me what's most interesting. Show me the key patterns, trends, and create 2-3 charts that best represent the data.",
                
                "📊 Show me charts": "Create 3 different charts from this data that show the most important insights. Use bar charts, line charts, or other appropriate visualizations.",
                
                "🔍 Find patterns": "Look for interesting patterns, trends, and relationships in this data. Highlight anything unusual or noteworthy.",
                
                "📈 Executive summary": "Create an executive summary of this data suitable for leadership, including key metrics, trends, and 1-2 supporting charts.",
                
                "🎯 Action insights": "Analyze this data and provide actionable insights and recommendations based on what the data shows.",
                
                "📑 Meeting summary": "Summarize this meeting transcript with key decisions, action items, and next steps."
            }
            
            for title, template in quick_templates.items():
                st.markdown(f"**{title}**")
                with st.container(border=True):
                    st.text_area("Template content:", value=template, height=80, key=f"quick_template_{title}", label_visibility="collapsed")
        
        # Usage instructions
        st.divider()
        st.success("💡 **How to use:** Select any template text above, copy it (Ctrl+C / Cmd+C), then paste it into the chat input below!")

    # Official Streamlit chat input - automatically positioned at bottom
    prompt = st.chat_input("Ask questions about your documents and data...", key="chat_input")
    
    # Bottom action buttons - positioned for fixed placement via CSS
    bottom_buttons_container = st.container(key="bottom_buttons_container")
    with bottom_buttons_container:
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            if st.button(type="primary", label="💡 Prompt Ideas", key="templates_bottom_button", use_container_width=True, help="Get AI-powered prompt suggestions based on your content"):
                show_prompt_assistant()
        
        with col2:
            if st.button("💬 Conversations", key="conversations_bottom_button", use_container_width=True, help="Browse and load saved conversations"):
                show_conversation_browser()
        
        with col3:
            if st.button("📊 Data Sources", key="data_sources_bottom_button", use_container_width=True, help="Manage your document and tabular data sources"):
                lib.show_data_sources()
        
        with col4:
            if st.button("⚙️ Model Settings", key="model_settings_bottom_button", use_container_width=True, help="Configure AI model and system settings"):
                show_model_settings()
    
    # Process the prompt when submitted
    if prompt:
        # Clear any pending confirmations when user sends a new message
        if "confirm_clear_chat" in st.session_state:
            st.session_state["confirm_clear_chat"] = False
            
        # Add user message to chat history
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        # Display user message immediately
        with st.chat_message("user"):
            st.markdown(prompt)
        
        # Generate response using UNIFIED AGENT
        with st.chat_message("assistant"):
            try:
                with st.spinner("Analyzing your request..."):
                    # Check if tool-calling is available
                    if hasattr(lib, 'TOOLS_AVAILABLE') and lib.TOOLS_AVAILABLE:
                        # Create unified agent
                        agent_executor = lib.create_unified_agent()
                        
                        if agent_executor is not None:
                            # ===== PHASE 1 CONTEXT OPTIMIZATION =====
                            # Get context statistics before processing
                            context_stats = lib.get_context_stats(st.session_state.messages)
                            
                            # Show context warning if approaching limits
                            if context_stats["estimated_tokens"] > 50000:  # Warning at ~50k tokens
                                st.info("📊 **Context Management**: Large conversation detected. Optimizing context to maintain performance...")
                            
                            # Execute the query using the unified agent
                            try:
                                result = agent_executor.invoke({"input": prompt})
                                
                                # ===== OPTIMIZE AGENT RESULT FOR CONTEXT =====
                                # Apply tool output optimization to save context space
                                # Use user-configurable tool output max length
                                tool_output_max = st.session_state.get("tool_output_max", 1000)
                                
                                # Customize the optimization for this session
                                if "intermediate_steps" in result:
                                    optimized_steps = []
                                    for action, observation in result["intermediate_steps"]:
                                        truncated_observation = lib.truncate_tool_output(observation, max_length=tool_output_max)
                                        optimized_steps.append((action, truncated_observation))
                                    
                                    result = result.copy()
                                    result["intermediate_steps"] = optimized_steps
                                
                                response = result["output"]
                                
                                # Check if the response indicates max iterations was reached
                                max_iterations = st.session_state.get("max_iterations", 10)
                                if "Agent stopped due to max iterations" in response or len(result.get("intermediate_steps", [])) >= max_iterations:
                                    st.warning("⚠️ The AI agent reached its maximum number of tool calls. The response may be incomplete.")
                                    response += "\n\n💡 **Tip**: Try asking a more specific question or break your request into smaller parts for better results."
                                
                            except Exception as agent_error:
                                # Handle specific agent errors
                                error_str = str(agent_error)
                                if "max iterations" in error_str.lower():
                                    response = """I apologize, but I reached the maximum number of analysis steps while processing your request. This usually happens with very complex queries.

**What you can do:**
1. **Break down your question** into smaller, more specific parts
2. **Be more specific** about which dataset or document you want me to analyze
3. **Try a simpler version** of your question first

**Examples of more specific questions:**
- Instead of "analyze everything", try "show me refugee trends for Q4"
- Instead of "create charts for all data", try "create a bar chart of funding by region"
- Instead of "what does the data show", try "what are the top 5 programs by effectiveness"

Please try rephrasing your question more specifically, and I'll be happy to help!"""
                                    st.warning("⚠️ Analysis complexity limit reached")
                                else:
                                    response = f"I encountered an error while processing your request: {error_str}\n\nPlease try rephrasing your question or contact support if the issue persists."
                                    st.error("An error occurred during analysis")
                                
                            # Display the response
                            st.markdown(response)
                            
                            # Check if charts were created and display them
                            chart_ids_to_display = []
                            chart_ids_for_response = ""
                            
                            # Check for multiple charts created in this response
                            if "current_chart_ids" in st.session_state and st.session_state.current_chart_ids:
                                chart_ids_to_display = st.session_state.current_chart_ids.copy()
                                if st.session_state.get("show_debug", False):
                                    st.write(f"🔧 DEBUG: Found current_chart_ids: {chart_ids_to_display}")
                            # Fallback to single chart ID for backwards compatibility
                            elif "current_chart_id" in st.session_state and st.session_state.current_chart_id is not None:
                                chart_ids_to_display = [st.session_state.current_chart_id]
                                if st.session_state.get("show_debug", False):
                                    st.write(f"🔧 DEBUG: Found single current_chart_id: {st.session_state.current_chart_id}")
                            
                            # Display all charts that were created
                            for chart_id in chart_ids_to_display:
                                if st.session_state.get("show_debug", False):
                                    st.write(f"🔧 DEBUG: Processing chart ID: {chart_id}")
                                    st.write(f"🔧 DEBUG: Stored charts keys: {list(st.session_state.get('stored_charts', {}).keys())}")
                                
                                if chart_id in st.session_state.get("stored_charts", {}):
                                    chart_config = st.session_state.stored_charts[chart_id]
                                    if st.session_state.get("show_debug", False):
                                        st.write(f"🔧 DEBUG: About to recreate chart with config: {chart_config}")
                                    
                                    fig = lib.recreate_chart_from_config(chart_config)
                                    if fig is not None:
                                        if st.session_state.get("show_debug", False):
                                            st.write(f"🔧 DEBUG: Chart recreation successful, displaying chart")
                                        
                                        # Anti-aliasing fix: Use specific config to reduce rendering artifacts
                                        config = {
                                            'displayModeBar': True,
                                            'displaylogo': False,
                                            'modeBarButtonsToRemove': ['pan2d', 'lasso2d', 'select2d'],
                                            'toImageButtonOptions': {
                                                'format': 'svg',  # SVG format for crisp rendering
                                                'filename': 'chart',
                                                'height': 500,
                                                'width': 700,
                                                'scale': 1  # No scaling to prevent artifacts
                                            }
                                        }
                                        st.plotly_chart(fig, use_container_width=True, key=f"current_chart_{chart_id}", config=config)
                                        
                                        # Collect chart IDs for adding to response
                                        chart_ids_for_response += f"\n\n[CHART_ID:{chart_id}]"
                                    else:
                                        # Show fallback message if chart recreation failed
                                        st.warning(f"⚠️ Chart {chart_id} could not be displayed. The chart data may be corrupted or incompatible with the current session.")
                                        if st.session_state.get("show_debug", False):
                                            st.write("**Debug Info:** Chart config:", chart_config)
                                else:
                                    if st.session_state.get("show_debug", False):
                                        st.write(f"🔧 DEBUG: Chart ID {chart_id} not found in stored_charts")
                            
                            # Add all chart IDs to response for persistence in chat history
                            if chart_ids_for_response:
                                response += chart_ids_for_response
                            
                            # Clear the current chart IDs after displaying
                            if "current_chart_ids" in st.session_state:
                                st.session_state.current_chart_ids = []
                            st.session_state.current_chart_id = None
                            
                            # Show debug information if enabled
                            if st.session_state.get("show_debug", False):
                                with st.expander("🔍 Agent Execution Details (Debug Mode)", expanded=False):
                                    # Show the tools that were used
                                    if "intermediate_steps" in result and result["intermediate_steps"]:
                                        st.write("### 🛠️ Tools Used:")
                                        
                                        for i, (action, observation) in enumerate(result["intermediate_steps"]):
                                            # Extract tool name and input
                                            tool_name = action.tool if hasattr(action, 'tool') else 'Unknown Tool'
                                            tool_input = action.tool_input if hasattr(action, 'tool_input') else {}
                                            
                                            # Display tool usage in a nice format
                                            st.write(f"**Step {i+1}: {tool_name}**")
                                            
                                            # Show tool input parameters
                                            if tool_input:
                                                with st.container():
                                                    st.write("*Input parameters:*")
                                                    for key, value in tool_input.items():
                                                        st.write(f"- {key}: `{value}`")
                                            
                                            # Show tool output (truncated for readability)
                                            if observation:
                                                with st.container():
                                                    st.write("*Tool output:*")
                                                    # Truncate long outputs
                                                    if len(str(observation)) > 500:
                                                        truncated_output = str(observation)[:500] + "... (truncated)"
                                                        st.text_area("", value=truncated_output, height=100, key=f"debug_output_{i}", disabled=True)
                                                    else:
                                                        st.text_area("", value=str(observation), height=100, key=f"debug_output_{i}", disabled=True)
                                            
                                            if i < len(result["intermediate_steps"]) - 1:
                                                st.divider()
                            
                            # Add assistant response to chat history
                            st.session_state.messages.append({"role": "assistant", "content": response})
                        else:
                            # Fallback if agent creation fails
                            error = "Agent creation failed. Please check your configuration."
                            st.error(error)
                            st.session_state.messages.append({"role": "assistant", "content": error})
                    else:
                        # Fallback to basic approach if tools not available
                        st.warning("Advanced tool-calling is not available. Using basic chat mode.")
                        
                        # Create basic chat response with context management
                        system_message = {"role": "system", "content": "You are a helpful assistant. Answer questions based on your knowledge."}
                        
                        # ===== APPLY CONTEXT WINDOW MANAGEMENT =====
                        all_session_messages = [
                            {"role": m["role"], "content": m["content"]} 
                            for m in st.session_state.messages
                        ]
                        
                        # Apply context optimization for long conversations using user settings
                        max_context_messages = st.session_state.get("max_context_messages", 20)
                        preserve_recent = st.session_state.get("preserve_recent", 10)
                        
                        optimized_messages = lib.manage_context_window(
                            all_session_messages, 
                            max_context_messages=max_context_messages, 
                            preserve_recent=preserve_recent
                        )
                        
                        # Combine system message with optimized messages
                        messages = [system_message] + optimized_messages
                        
                        # Create streaming response
                        stream = client.chat.completions.create(
                            model=st.session_state["openai_model"],
                            messages=messages,
                            stream=True,
                        )
                        
                        # Display streaming response
                        response = st.write_stream(stream)
                        
                        # Add assistant response to chat history
                        st.session_state.messages.append({"role": "assistant", "content": response})
                        
            except openai.APIConnectionError as e:
                error = "Sorry, the server could not be reached. Please try again later..."
                st.error(error)
                st.session_state.messages.append({"role": "assistant", "content": error})
                
            except openai.RateLimitError as e:
                error = "Too many requests. Please try again later..."
                st.error(error)
                st.session_state.messages.append({"role": "assistant", "content": error})
                
            except openai.APIStatusError as e:
                error = "An error occurred while processing your request. Please try again later..."
                st.error(error)
                st.session_state.messages.append({"role": "assistant", "content": error})
                
            except Exception as e:
                error = f"An unexpected error occurred: {str(e)}"
                st.error(error)
                st.session_state.messages.append({"role": "assistant", "content": error})
        
        # Trigger auto-save after processing response
        auto_save_conversation()
        
        # Apply automatic context trimming if conversation is getting too large
        # This prevents gradual token accumulation over long sessions
        if len(st.session_state.messages) > 40:
            context_stats = lib.get_context_stats(st.session_state.messages)
            
            # Get current model's limit
            current_model = st.session_state.get("openai_model", "gpt-3.5-turbo")
            limit_info = lib.get_model_token_limit(current_model)
            token_limit = limit_info["limit"]
            
            # If we're at 80% of the limit, start trimming proactively
            if context_stats["estimated_tokens"] > (token_limit * 0.8):
                original_count = len(st.session_state.messages)
                st.session_state.messages = lib.manage_context_window(
                    st.session_state.messages,
                    max_context_messages=30,
                    preserve_recent=20
                )
                trimmed_count = original_count - len(st.session_state.messages)
                
                if trimmed_count > 0 and st.session_state.get("show_debug", False):
                    st.info(f"🔄 Auto-trimmed {trimmed_count} older messages to maintain optimal performance.")

    # Conversation Management - Always visible at bottom when there are messages
    if len(st.session_state.messages) > 0:
        st.markdown("---")
        st.markdown("### 📋 Conversation Management")
        
        # Use 5 columns for compact buttons
        col1, col2, col3, col4, col5 = st.columns([1, 1, 1, 1, 1])
        
        with col1:
            # Manual save conversation button
            if st.button("💾 Save", use_container_width=True, help="Manually save conversation"):
                manual_save_conversation()
        
        with col2:
            # Download answers as DOCX
            doc_download = Document()
            if "saved_anonymisation" in st.session_state:
                doc_title = st.session_state["saved_anonymisation"]["Title"].replace(">"," ")
            else:
                doc_title = "Chat Conversation"
            # Filter the title for invalid XML characters
            doc_title = lib.filter_xml_chars(doc_title)
            doc_download.add_heading(doc_title, level=1)
            
            # Add conversation to document
            for message in st.session_state.messages:
                if message["role"] == "user":
                    question_text = lib.filter_xml_chars(message['content'])
                    doc_download.add_heading(f"Question: {question_text}", level=2)
                else:
                    answer_text = lib.filter_xml_chars(message['content'])
                    doc_download.add_paragraph(answer_text)
            
            # Save document to buffer
            bio = io.BytesIO()
            doc_download.save(bio)
            
            # Download button
            st.download_button(
                label="📥 Download",
                type="secondary",
                data=bio.getvalue(),
                file_name=f"ChatGPT Answers About {doc_title}.docx",
                mime="docx",
                key="download_conversation_btn",
                use_container_width=True,
                help="Download conversation as DOCX"
            )
        
        with col3:
            # Copy answers to clipboard
            copy_clicked = st.button("📋 Copy", use_container_width=True, help="Copy to clipboard")
        
        with col4:
            # Clear conversation with confirmation
            clear_clicked = st.button("🗑️ Clear", use_container_width=True, help="Clear chat history")
        
        with col5:
            # Save and proceed to reverse anonymization
            if "saved_anonymisation" in st.session_state:
                entities = st.session_state["saved_anonymisation"]["Entities"]
                if "saved_content" in st.session_state:
                    attendees = st.session_state["saved_content"]["Attendees"]
                else:
                    attendees = {}
                
                # Format conversation for saving
                conversation_text = ""
                for message in st.session_state.messages:
                    if message["role"] == "user":
                        conversation_text += f"Question: {message['content']}\n\n"
                    else:
                        conversation_text += f"Answer: {message['content']}\n\n"
                
                # Save button
                if st.button(
                    label="Save & Reverse",
                    type="primary",
                    use_container_width=True,
                    key="save_conversation",
                    on_click=lib.save_gpt_answers,
                    args=[st.session_state["saved_anonymisation"]["Title"], conversation_text, entities, attendees]
                ):
                    st.switch_page("pages/03_revert.py")
            else:
                # If no saved_anonymisation, show disabled button with tooltip
                st.button(
                    label="Save & Reverse",
                    type="primary",
                    disabled=True,
                    use_container_width=True,
                    key="save_conversation_disabled",
                    help="Process a document first to enable this feature"
                )
        
        # Handle button actions and display messages outside column layout
        if copy_clicked:
            # Clear any pending confirmations when user performs other actions
            if "confirm_clear_chat" in st.session_state:
                st.session_state["confirm_clear_chat"] = False
                
            # Format conversation for clipboard
            conversation_text = ""
            for message in st.session_state.messages:
                if message["role"] == "user":
                    conversation_text += f"Question: {message['content']}\n\n"
                else:
                    conversation_text += f"Answer: {message['content']}\n\n"
            
            # Copy to clipboard
            pyperclip.copy(conversation_text)
            st.success("✅ Conversation copied to clipboard!")
        
        if clear_clicked:
            if st.session_state.get("confirm_clear_chat"):
                # If already confirmed, perform the deletion
                st.session_state.messages = []
                # Also clear stored charts to free up memory
                if "stored_charts" in st.session_state:
                    st.session_state.stored_charts = {}
                if "current_chart_id" in st.session_state:
                    st.session_state.current_chart_id = None
                if "current_chart_ids" in st.session_state:
                    st.session_state.current_chart_ids = []
                # Clear conversation tracking
                st.session_state.current_conversation_id = None
                st.session_state.last_conversation_save_time = None
                st.session_state.last_save_message_count = 0
                # Clear confirmation state
                st.session_state["confirm_clear_chat"] = False
                st.success("✅ Chat history cleared successfully!")
                st.rerun()
            else:
                # Check for unsaved changes
                save_status = lib.get_conversation_save_status()
                
                # Set confirmation state and show appropriate warning
                st.session_state["confirm_clear_chat"] = True
                
                if save_status["has_unsaved"]:
                    unsaved_count = save_status["message_count"] - save_status["last_save_message_count"]
                    st.error(f"⚠️ **WARNING:** You have {unsaved_count} unsaved message(s)! Click '🗑️ Clear' again to confirm deletion.")
                    st.info("💡 **Tip:** Use the '💾 Save' button to save your conversation before clearing.")
                else:
                    st.warning("⚠️ Click '🗑️ Clear' again to confirm. This action cannot be undone!")
        
        # Show any pending confirmation messages
        if st.session_state.get("confirm_clear_chat") and not clear_clicked:
            save_status = lib.get_conversation_save_status()
            if save_status["has_unsaved"]:
                unsaved_count = save_status["message_count"] - save_status["last_save_message_count"]
                st.error(f"⚠️ **Confirmation pending:** {unsaved_count} unsaved message(s) will be lost! Click '🗑️ Clear' again to confirm.")
            else:
                st.warning("⚠️ Confirmation pending: Click '🗑️ Clear' again to confirm deletion.")

def display_tabular_sources():
    """Display available tabular data sources with management options"""
    available_datasets = st.session_state.get("tabular_datasets", {})
    
    if available_datasets:
        st.subheader("📊 Tabular Data Sources")
        
        # Add debug mode toggle
        col1, col2 = st.columns([3, 1])
        with col1:
            st.write(f"**{len(available_datasets)} dataset(s) available for analysis:**")
        with col2:
            debug_mode = st.checkbox("🔧 Debug Mode", 
                                   value=st.session_state.get("show_debug", False),
                                   help="Show detailed debugging information for chart generation")
            st.session_state.show_debug = debug_mode
        
        # Display datasets in columns
        cols = st.columns(min(len(available_datasets), 3))
        for idx, (name, df) in enumerate(available_datasets.items()):
            with cols[idx % 3]:
                with st.expander(f"📊 {name}", expanded=False):
                    st.write(f"**Rows:** {df.shape[0]:,}")
                    st.write(f"**Columns:** {df.shape[1]}")
                    st.write("**Column names:**")
                    for col in df.columns:
                        st.write(f"• {col}")
                    
                    # Show sample data
                    st.write("**Sample data:**")
                    st.dataframe(df.head(3), use_container_width=True)
                    
                    # Add data type info in debug mode
                    if st.session_state.get("show_debug", False):
                        st.write("**Data types:**")
                        for col, dtype in df.dtypes.items():
                            st.write(f"• {col}: {dtype}")
        
        st.markdown("---")
        st.markdown("💡 **Tip:** Reference column names exactly as shown above when requesting specific charts.")
        
        # Clear data option
        if st.button("🗑️ Clear All Tabular Data", type="secondary"):
            st.session_state.tabular_datasets = {}
            st.success("All tabular data cleared!")
            st.rerun()
    else:
        st.info("📊 No tabular datasets loaded. Upload CSV or Excel files on the Home page to get started.")
