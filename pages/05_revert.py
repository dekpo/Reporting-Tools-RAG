# Packages
import streamlit as st
from docx import Document
import pyperclip

# Modules
import io

# Utils
import lib

# The App config and styling
lib.app_config()

# Sidebar and Navbar
lib.sidebar()

st.title(lib.APP_TITLE)

st.divider()

# Steps
lib.steps(4)

st.divider()

st.header("Reverse Anonymization")

st.markdown("<p>Re-insert entities ( people, organizations, locations...) to your contents using previously saved data.</p>",unsafe_allow_html=True)

st.subheader("Your Content's Entities")
if "saved_gpt_answers" in st.session_state:
    saved_entities = st.session_state["saved_gpt_answers"]["Entities"]
    
    # Check if entities have the expected structure
    if not saved_entities or "Text" not in saved_entities:
        st.warning("The entities data is not in the expected format. This might be because you're accessing a document from a previous session.")
        
        # Create a default structure for entities if missing
        saved_entities = {
            "Text": [],
            "Replacement": [],
            "Category": []
        }
        # Update the session state with the default structure
        st.session_state["saved_gpt_answers"]["Entities"] = saved_entities
    
    st.write(f'You have **{len(saved_entities["Text"])}** entities to reverse.')

    edited_entities = st.data_editor(saved_entities,disabled=["Replacement","Category"],use_container_width=True)
    
    saved_attendees = st.session_state["saved_gpt_answers"]["Attendees"]

    # Check if attendees have the expected structure
    if not saved_attendees or not isinstance(saved_attendees, dict):
        saved_attendees = {}
        st.session_state["saved_gpt_answers"]["Attendees"] = saved_attendees
    elif "Attendee" not in saved_attendees and len(saved_attendees) > 0:
        st.warning("The attendees data is not in the expected format.")
        saved_attendees = {
            "Attendee": [],
            "Replacement": []
        }
        st.session_state["saved_gpt_answers"]["Attendees"] = saved_attendees

    # Initialize edited_attendees to None
    edited_attendees = None
    
    if saved_attendees and "Attendee" in saved_attendees and len(saved_attendees["Attendee"]) > 0:
        st.subheader("The Meeting Attendees")
        st.write(f'You have **{len(saved_attendees["Attendee"])}** attendees to re-insert.')

        edited_attendees = st.data_editor(saved_attendees,disabled=["Replacement"],use_container_width=True)
    
    reverse_now = st.checkbox('**Re-insert All Entities Now ?**',key="reverse_now",value=True,help="This feature can automatically re-insert entities in the content based on your backup.")

    st.divider()
    
    st.write("""<div id='top-content'></div>""",unsafe_allow_html=True)

    st.subheader("Your Anonymized Content To Reverse is Below:")

    st.markdown("<p><a href='#download-or-save-your-data'>🔽Download Or Copy Your Data At The Bottom Of This Page🔽</a></p>",unsafe_allow_html=True)

    content = st.container(border=True)

    content.write("""<div class='extracted-content' />""",unsafe_allow_html=True)

    txt = st.session_state["saved_gpt_answers"]["Data"]

    if reverse_now:
        # Check if entities have the expected structure before processing
        if "Replacement" in edited_entities and "Text" in edited_entities and "Category" in edited_entities:
            for term in edited_entities["Replacement"]:
                if term:  # Only process non-empty terms
                    index = edited_entities["Replacement"].index(term)
                    if index < len(edited_entities["Text"]) and index < len(edited_entities["Category"]):
                        replacement = edited_entities["Text"][index]
                        category = edited_entities["Category"][index]
                        txt = txt.replace(term, f'<span class="entity {category.lower()}">{replacement}</span>')
        
        # Check if attendees have the expected structure before processing
        if saved_attendees and "Attendee" in saved_attendees and "Replacement" in saved_attendees:
            for i, attendee in enumerate(saved_attendees["Replacement"]):
                if attendee and i < len(saved_attendees["Attendee"]):
                    replacement = saved_attendees["Attendee"][i]
                    txt = txt.replace(attendee, f'<span class="entity person">{replacement}</span>')
    
    txt = txt.replace("$","\\$")
    txt = txt.replace("\n","<br>")
    content.markdown(txt,unsafe_allow_html=True)
    txt = txt.replace("<br>","\n")
    raw_text = lib.strip_tags(txt)

    st.divider()

    st.subheader("Download Or Save Your Data")
            
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button(f"Copy Your Content To Clipboard",use_container_width=True):
            gpt_answers = raw_text
            pyperclip.copy(gpt_answers)
            st.success("Answers Copied !")
    with col2:
        st.write("OR")
    with col3:
        doc_download = Document()
        doc_title = st.session_state["saved_gpt_answers"]["Title"].replace(">"," ")
        doc_download.add_heading(doc_title,level=1)
        doc_download.add_paragraph(raw_text)
        bio = io.BytesIO()
        doc_download.save(bio)
        if doc_download:
            st.download_button(
            label="**Download Your Content As DOCX File**",
            data=bio.getvalue(),
            type="primary",
            file_name=f"ChatGPT Answers About {doc_title}.docx",
            mime="docx"
            )
    st.divider()
    st.markdown("<p><a href='#top-content'>🔼Go Back To The Top Of This Content🔼</a></p>",unsafe_allow_html=True)
else:
    st.markdown('<p>You have <b>no entity</b> to re-insert. You must first apply anonymization with the previous tool to be able to keep references of your entities.</p>',unsafe_allow_html=True)
    st.page_link(page="pages/03_anonymize.py",label="Got To Anonymize Content",icon=":material/sms:",use_container_width=True)
#st.write(st.session_state)