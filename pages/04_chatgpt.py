# Packages
import streamlit as st
import openai
from openai import OpenAI
from docx import Document
import pyperclip

# Modules
import time
import io
import os
from datetime import datetime

# Utils
import lib

# The App config and styling
lib.app_config()

# Sidebar and Navbar
lib.sidebar()

st.title(lib.APP_TITLE)

st.divider()

# Steps
lib.steps(3)

st.divider()

if "gpt_api_key" not in st.session_state:
    st.header("ChatGPT Tool")

    st.markdown("<p>Ask ChatGPT to answer questions about your anonymized content using advanced retrieval techniques.</p>",unsafe_allow_html=True)

    # Check if we have a persistent ChromaDB to load
    persist_directory = './chroma_db'
    if os.path.exists(persist_directory) and os.path.isdir(persist_directory):
        document_metadata = lib.load_document_metadata(persist_directory)
        if document_metadata:
            st.info(f"Found {len(document_metadata)} document(s) in storage. After submitting your API key, you'll be able to access them.")

    MY_API_KEY = st.text_input(
            label="**Please specify your OpenAI API Key** this one is for testing purpose.",
            value=st.secrets["OPENAI_API_KEY"]
        )

    st.markdown("<a href=\"https://platform.openai.com/api-keys\" target=\"_blank\">Get your OpenAI API key here !</a>",unsafe_allow_html=True)

    # Define model options with categorization
    model_options = {
        "GPT-4 Models (Most Capable)": [
            "gpt-4o", 
            "gpt-4-turbo", 
            "gpt-4"
        ],
        "GPT-3.5 Models (Fast & Cost-effective)": [
            "gpt-3.5-turbo"
        ]
    }

    # Flatten the options for selectbox
    all_models = []
    for category, models in model_options.items():
        for model in models:
            all_models.append(f"{category}: {model}")

    # Model descriptions for information
    model_descriptions = {
        "gpt-4o": "Latest and most capable model, optimized for performance and cost",
        "gpt-4-turbo": "Powerful model with strong reasoning capabilities",
        "gpt-4": "Original GPT-4 model with high accuracy",
        "gpt-3.5-turbo": "Fast and cost-effective model for most general tasks"
    }

    # Select model with categorization
    selected_model_with_category = st.selectbox(
        "**Select GPT model**",
        all_models,
        index=0,
        help="Choose the OpenAI model you want to use. More capable models provide better results but may cost more."
    )

    # Extract just the model name
    selected_model = selected_model_with_category.split(": ")[-1]
    
    # Show model description if available
    if selected_model in model_descriptions:
        st.caption(model_descriptions[selected_model])

    # Store in session state
    if "openai_model" not in st.session_state:
        st.session_state["openai_model"] = selected_model
    elif st.session_state["openai_model"] != selected_model:
        st.session_state["openai_model"] = selected_model
    

    if st.button("Submit Your API Key",key="submit_your_api_key_btn1",type="primary"):
        client = OpenAI(api_key=MY_API_KEY)
        try:
            # Check the validity of the API key with the selected model
            try:
                # First check if the model exists and is available
                models = client.models.list()
                model_available = False
                available_models = [model.id for model in models.data]
                
                # Check if the exact model is available
                if selected_model in available_models:
                    model_available = True
                # Some models might have different names in the API (e.g., gpt-4o might be available as gpt-4o-xxxx)
                elif any(model_id.startswith(selected_model) for model_id in available_models):
                    model_available = True
                    
                if not model_available:
                    # If model not available, suggest a fallback
                    fallback_model = "gpt-3.5-turbo"  # Default fallback
                    st.warning(f"The selected model '{selected_model}' appears to be unavailable with your API key. Falling back to {fallback_model}.", icon="⚠️")
                    selected_model = fallback_model
                    st.session_state["openai_model"] = selected_model
                
                # Now try to use the model
                client.chat.completions.create(
                    model=selected_model,
                    messages=[
                        {"role": "user", "content": "Hello this is a test."}
                    ],
                    stream=False,
                )
            except Exception as model_error:
                st.warning(f"Error with model {selected_model}: {model_error}. Trying with gpt-3.5-turbo instead.", icon="⚠️")
                # Fallback to gpt-3.5-turbo which is widely available
                client.chat.completions.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "user", "content": "Hello this is a test."}
                    ],
                    stream=False,
                )
                # Update the model to the fallback
                st.session_state["openai_model"] = "gpt-3.5-turbo"
                
        except Exception as e:
            st.error(f"Your API key is not valid. Please try Again. Error: {str(e)}", icon="⚠️")
            if 'submit_your_api_key_btn' in st.session_state:
                del st.session_state.submit_your_api_key_btn
        else:
            st.success(f"Your API key is valid. You will be redirected to the discussion area in few seconds. Using model: {st.session_state['openai_model']}", icon="✅")
            st.session_state["gpt_api_key"] = MY_API_KEY
            
            # Check for existing documents in the persistent storage
            persist_directory = './chroma_db'
            if os.path.exists(persist_directory) and os.path.isdir(persist_directory):
                document_metadata = lib.load_document_metadata(persist_directory)
                if document_metadata:
                    with st.spinner("Loading document database..."):
                        try:
                            # Initialize ChromaDB client with custom embedding function
                            if hasattr(lib, 'RAG_AVAILABLE') and lib.RAG_AVAILABLE:
                                # Create custom embedding function
                                embedding_function = lib.OpenAIEmbeddingFunction(api_key=MY_API_KEY)
                                
                                # Create ChromaDB client
                                import chromadb
                                from chromadb.config import Settings
                                
                                chroma_client = chromadb.PersistentClient(
                                    path=persist_directory,
                                    settings=Settings(
                                        anonymized_telemetry=False,
                                        allow_reset=True
                                    )
                                )
                                
                                # Connect to the collection
                                collection = chroma_client.get_collection(
                                    name="document_collection", 
                                    embedding_function=embedding_function
                                )
                                
                                # Store in session state
                                st.session_state.vector_db = collection
                                st.session_state.selected_doc_sources = list(document_metadata.keys())
                                
                                # If there are documents available, select the first one as the active document
                                # This allows users to immediately use their saved documents
                                if document_metadata and "saved_anonymisation" not in st.session_state:
                                    # Get the most recent document (sorted by timestamp)
                                    doc_list = []
                                    for doc_hash, doc_data in document_metadata.items():
                                        doc_list.append({
                                            "hash": doc_hash,
                                            "title": doc_data["title"],
                                            "timestamp": doc_data.get("timestamp", 0)
                                        })
                                    # Sort by timestamp (newest first)
                                    doc_list.sort(key=lambda x: x["timestamp"], reverse=True)
                                    
                                    if doc_list:
                                        # Set the latest document as the active document
                                        latest_doc = doc_list[0]
                                        st.session_state["saved_anonymisation"] = {
                                            "Title": latest_doc["title"],
                                            "Time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                            "Entities": {}  # Entities might not be necessary for querying
                                        }
                        except Exception as e:
                            st.error(f"Error loading document database: {e}")
            
            time.sleep(1)
            st.rerun()
else:
    st.header(f"ChatGPT Discussion (model: {st.session_state["openai_model"] })")
    if st.button("Reset Your API Key And GPT model",key="reset_your_api_key_btn"):
        del st.session_state["gpt_api_key"]
        del st.session_state["openai_model"]
        # Clear vector DB session state if it exists
        if "vector_db" in st.session_state:
            del st.session_state["vector_db"]
        if "processed_anonymizations" in st.session_state:
            del st.session_state["processed_anonymizations"]
        st.rerun()

    # Document Source Selection - Moved above the chat interface
    if "vector_db" in st.session_state and st.session_state.vector_db is not None:
        st.divider()
        st.subheader("Document Sources")
        
        # Load document metadata
        persist_directory = './chroma_db'
        document_metadata = lib.load_document_metadata(persist_directory)
        
        if document_metadata:
            # Display document list with selection options
            st.write("Select documents to use as sources for your questions:")
            
            # Convert metadata to a more usable format for display
            doc_list = []
            for doc_hash, doc_data in document_metadata.items():
                doc_list.append({
                    "hash": doc_hash,
                    "title": doc_data["title"],
                    "chunks": doc_data["chunks"],
                    "timestamp": doc_data.get("timestamp", 0),
                    "active": doc_data.get("active", True)
                })
            
            # Sort by timestamp (newest first)
            doc_list.sort(key=lambda x: x["timestamp"], reverse=True)
            
            # Create selection interface
            selected_docs = []
            for doc in doc_list:
                col1, col2, col3 = st.columns([0.6, 0.2, 0.2])
                with col1:
                    selected = st.checkbox(
                        f"{doc['title']} ({doc['chunks']} chunks)", 
                        value=doc['active'],
                        key=f"doc_select_{doc['hash']}"
                    )
                    if selected:
                        selected_docs.append(doc['hash'])
                
                with col2:
                    timestamp = datetime.fromtimestamp(doc["timestamp"])
                    st.caption(f"Added: {timestamp.strftime('%Y-%m-%d')}")
                
                with col3:
                    if st.button("Remove", key=f"remove_{doc['hash']}"):
                        if st.session_state.get("confirm_delete") == doc['hash']:
                            # If already confirmed, perform the deletion
                            with st.spinner(f"Deleting '{doc['title']}'..."):
                                success = lib.delete_document_from_vector_db(doc['hash'])
                                if success:
                                    st.rerun()
                                # Clear confirmation state
                                st.session_state["confirm_delete"] = None
                        else:
                            # Set confirmation state and show confirmation message
                            st.session_state["confirm_delete"] = doc['hash']
                            st.warning("Click 'Remove' again to confirm deletion.", icon="⚠️")
            
            # Store selected documents in session state
            if "selected_doc_sources" not in st.session_state or st.session_state.selected_doc_sources != selected_docs:
                st.session_state.selected_doc_sources = selected_docs
            
            # Update metadata with active status
            for doc_hash in document_metadata:
                document_metadata[doc_hash]['active'] = doc_hash in selected_docs
            lib.save_document_metadata(persist_directory, document_metadata)
        else:
            st.info("No documents found in the database. Process a document to add it to the sources.")
        st.divider()

    st.markdown("<p>Ask questions about your anonymized content. The AI will provide relevant answers based on the context of your document.</p>",unsafe_allow_html=True)

    client = OpenAI(api_key=st.session_state["gpt_api_key"])

    # Initialize messages in session state if they don't exist
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # Initialize processed_anonymizations if it doesn't exist
    if "processed_anonymizations" not in st.session_state:
        st.session_state.processed_anonymizations = set()

    # Process document and create vector database if not already processed
    if "saved_anonymisation" in st.session_state:
        # Generate a unique key for this anonymization
        anon_key = st.session_state["saved_anonymisation"]["Title"] + "_" + st.session_state["saved_anonymisation"]["Time"]
        
        # Check if this anonymization has already been processed
        if anon_key not in st.session_state.processed_anonymizations and "vector_db" in st.session_state:
            # We have a new document to add to the existing vector DB
            # Only process if we have actual document data
            if "Data" in st.session_state["saved_anonymisation"] and st.session_state["saved_anonymisation"]["Data"]:
                with st.spinner(f"Adding new document '{st.session_state['saved_anonymisation']['Title']}' to the database..."):
                    # Get document info
                    document_title = st.session_state["saved_anonymisation"]["Title"]
                    document_content = st.session_state["saved_anonymisation"]["Data"]
                    
                    # Add to vector database
                    st.session_state.vector_db = lib.create_vector_db_from_text(
                        document_content, 
                        document_title, 
                        st.session_state["gpt_api_key"]
                    )
                    
                    if st.session_state.vector_db is not None:
                        st.success(f"Document '{document_title}' added successfully to the database!")
                        # Mark as processed
                        st.session_state.processed_anonymizations.add(anon_key)
                    else:
                        st.error(f"Failed to add document '{document_title}' to the database.")
                
        elif anon_key not in st.session_state.processed_anonymizations and "vector_db" not in st.session_state:
            # This is the first document, and we need to create the vector DB
            # Only process if we have actual document data
            if "Data" in st.session_state["saved_anonymisation"] and st.session_state["saved_anonymisation"]["Data"]:
                with st.spinner("Processing your document... This might take a minute."):
                    # Check if RAG is available
                    if not hasattr(lib, 'RAG_AVAILABLE') or not lib.RAG_AVAILABLE:
                        st.warning("RAG functionality is not available. The app will use the traditional chunking approach instead.")
                        # Set vector_db to None to skip RAG-related code
                        st.session_state.vector_db = None
                    else:
                        # Create vector database from anonymized content
                        document_title = st.session_state["saved_anonymisation"]["Title"]
                        document_content = st.session_state["saved_anonymisation"]["Data"]
                        
                        # Create vector database
                        st.session_state.vector_db = lib.create_vector_db_from_text(
                            document_content, 
                            document_title, 
                            st.session_state["gpt_api_key"]
                        )
                        
                        if st.session_state.vector_db is not None:
                            st.success(f"Document processed successfully! You can now ask questions about '{document_title}'")
                            # Mark as processed
                            st.session_state.processed_anonymizations.add(anon_key)
                        else:
                            st.error("Failed to process document with RAG. Falling back to traditional approach.")
                            # Set vector_db to None to indicate RAG is not available
                            st.session_state.vector_db = None

    # Display chat header and existing messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Create a container for document processing messages
    header = st.container()
    # header.write(f"<div class='fixed-header'/>",unsafe_allow_html=True)

    # Show a warning if no document is loaded AND no vector DB is available
    if "saved_anonymisation" not in st.session_state and ("vector_db" not in st.session_state or st.session_state.vector_db is None):
        st.warning("No document loaded. Please go to the Anonymize page to process a document first.")
    else:
        # Display a message based on whether RAG is available
        if "vector_db" not in st.session_state or st.session_state.vector_db is None:
            if hasattr(lib, 'RAG_AVAILABLE') and lib.RAG_AVAILABLE:
                st.warning("Document database not created. Please try processing your document again.")
            else:
                st.info("Using traditional question answering. RAG functionality is not available.")
        else:
            # Check if we're using a previously stored document without its content
            if "saved_anonymisation" in st.session_state and ("Data" not in st.session_state["saved_anonymisation"] or not st.session_state["saved_anonymisation"].get("Data")):
                st.info("You're viewing previously stored document sources. You can ask questions about these documents from previous sessions.")
        
        # Chat input for user questions
        if prompt := st.chat_input("Ask a question about your document", max_chars=8000):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            # Display user message
            with st.chat_message("user"):
                st.markdown(prompt)
            
            # Generate response using RAG
            with st.chat_message("assistant"):
                try:
                    with st.spinner("Searching document and generating response..."):
                        # Check if we can use RAG approach
                        if hasattr(lib, 'RAG_AVAILABLE') and lib.RAG_AVAILABLE and st.session_state.vector_db is not None:
                            # Query the vector database
                            context_docs, context_metadatas = lib.query_vector_db(
                                st.session_state.vector_db, 
                                prompt,
                                n_results=5,
                                selected_doc_sources=st.session_state.get("selected_doc_sources", [])
                            )
                            
                            # Log debugging information
                            st.session_state["last_query_debug"] = {
                                "query": prompt,
                                "found_contexts": len(context_docs),
                                "selected_sources": st.session_state.get("selected_doc_sources", [])
                            }
                            
                            if not context_docs:
                                # No relevant context found
                                response = "I couldn't find relevant information in the document to answer your question. Please try rephrasing your question or ask something else about the document content."
                                st.markdown(response)
                                st.session_state.messages.append({"role": "assistant", "content": response})
                            else:
                                # Debug info in dev environment
                                if "show_debug" in st.session_state and st.session_state.show_debug:
                                    with st.expander("Debug Information (Admin only)"):
                                        st.write(f"Found {len(context_docs)} context chunks")
                                        st.write("First few words of each chunk:")
                                        for i, doc in enumerate(context_docs[:3]):
                                            st.write(f"{i+1}. {doc[:50]}...")
                                
                                # Stream the response for a better user experience
                                stream = client.chat.completions.create(
                                    model=st.session_state["openai_model"],
                                    messages=[
                                        {"role": "system", "content": "You are a helpful assistant providing information based on the supplied document context. Answer questions accurately using ONLY information from the provided context. When referencing information, mention which document source it came from. If multiple documents contain relevant information, clearly indicate which source each piece of information came from. If the context doesn't contain information to answer the question, admit you don't know rather than making up an answer."},
                                        {"role": "user", "content": f"Context from {len(context_docs)} document chunks:\n\n{' '.join(context_docs)}\n\nSource metadata:\n{context_metadatas}\n\nUser Question: {prompt}\n\nImportant: Base your answer ONLY on the provided context. Cite document sources when possible."}
                                    ],
                                    stream=True,
                                )
                                
                                # Display streaming response
                                response = st.write_stream(stream)
                                
                                # Add a section showing the sources used
                                if context_metadatas:
                                    with st.expander("View document sources used"):
                                        sources_used = {}
                                        for metadata in context_metadatas:
                                            if "source" in metadata:
                                                source = metadata["source"]
                                                if source not in sources_used:
                                                    sources_used[source] = 0
                                                sources_used[source] += 1
                                        
                                        st.markdown("### Document sources used in this response:")
                                        for source, count in sorted(sources_used.items()):
                                            st.markdown(f"- **{source}** ({count} chunks)")
                                        
                                        if st.session_state.show_debug:
                                            st.divider()
                                            st.write("### Query Debug Info")
                                            if "last_query_debug" in st.session_state:
                                                debug = st.session_state["last_query_debug"]
                                                st.write(f"Query: '{debug['query']}'")
                                                st.write(f"Total chunks found: {debug['found_contexts']}")
                                                st.write(f"Selected sources: {', '.join(debug['selected_sources']) if debug['selected_sources'] else 'All sources'}")
                                
                                # Add assistant response to chat history
                                st.session_state.messages.append({"role": "assistant", "content": response})
                        else:
                            # Fallback to traditional approach when RAG is not available
                            # Stream the response for a better user experience
                            # Create messages list with system message and all existing messages
                            messages = [
                                {"role": "system", "content": "You are a helpful assistant. Answer questions based on your knowledge."}
                            ]
                            # Add all previous messages
                            messages.extend([
                                {"role": m["role"], "content": m["content"]} 
                                for m in st.session_state.messages
                            ])
                            
                            # Create streaming response
                            stream = client.chat.completions.create(
                                model=st.session_state["openai_model"],
                                messages=messages,
                                stream=True,
                            )
                            
                            # Display streaming response
                            response = st.write_stream(stream)
                            
                            # Add assistant response to chat history
                            st.session_state.messages.append({"role": "assistant", "content": response})
                            
                except openai.APIConnectionError as e:
                    error = "Sorry, the server could not be reached. Please try again later..."
                    st.error(error)
                    st.session_state.messages.append({"role": "assistant", "content": error})
                    
                except openai.RateLimitError as e:
                    error = "Too many requests. Please try again later..."
                    st.error(error)
                    st.session_state.messages.append({"role": "assistant", "content": error})
                    
                except openai.APIStatusError as e:
                    error = "An error occurred while processing your request. Please try again later..."
                    st.error(error)
                    st.session_state.messages.append({"role": "assistant", "content": error})
                    
                except Exception as e:
                    error = f"An unexpected error occurred: {str(e)}"
                    st.error(error)
                    st.session_state.messages.append({"role": "assistant", "content": error})

    # Document management interface
    if "saved_anonymisation" in st.session_state:
        st.divider()
        st.subheader("Document Management")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            # Download answers as DOCX
            if len(st.session_state.messages) > 0:
                doc_download = Document()
                doc_title = st.session_state["saved_anonymisation"]["Title"].replace(">"," ")
                doc_download.add_heading(doc_title, level=1)
                
                # Add conversation to document
                for message in st.session_state.messages:
                    if message["role"] == "user":
                        doc_download.add_heading(f"Question: {message['content']}", level=2)
                    else:
                        doc_download.add_paragraph(message["content"])
                
                # Save document to buffer
                bio = io.BytesIO()
                doc_download.save(bio)
                
                # Download button
                st.download_button(
                    label="Download Conversation As DOCX File",
                    type="secondary",
                    data=bio.getvalue(),
                    file_name=f"ChatGPT Answers About {doc_title}.docx",
                    mime="docx",
                    key="download_conversation_btn",
                    use_container_width=True
                )
        
        with col2:
            # Copy answers to clipboard
            if len(st.session_state.messages) > 0:
                if st.button("Copy Conversation To Clipboard", use_container_width=True):
                    # Format conversation for clipboard
                    conversation_text = ""
                    for message in st.session_state.messages:
                        if message["role"] == "user":
                            conversation_text += f"Question: {message['content']}\n\n"
                        else:
                            conversation_text += f"Answer: {message['content']}\n\n"
                    
                    # Copy to clipboard
                    pyperclip.copy(conversation_text)
                    st.success("Conversation Copied!")
            
            # Clear conversation
            if len(st.session_state.messages) > 0:
                if st.button("Clear Conversation", use_container_width=True):
                    st.session_state.messages = []
                    st.rerun()
        
        with col3:
            # Save and proceed to reverse anonymization
            if len(st.session_state.messages) > 0:
                entities = st.session_state["saved_anonymisation"]["Entities"]
                if "saved_content" in st.session_state:
                    attendees = st.session_state["saved_content"]["Attendees"]
                else:
                    attendees = {}
                
                # Format conversation for saving
                conversation_text = ""
                for message in st.session_state.messages:
                    if message["role"] == "user":
                        conversation_text += f"Question: {message['content']}\n\n"
                    else:
                        conversation_text += f"Answer: {message['content']}\n\n"
                
                # Save button
                if st.button(
                    label="**Save Conversation And >> Reverse Anonymization**",
                    type="primary",
                    use_container_width=True,
                    key="save_conversation",
                    on_click=lib.save_gpt_answers,
                    args=[st.session_state["saved_anonymisation"]["Title"], conversation_text, entities, attendees]
                ):
                    st.switch_page("pages/05_revert.py")
        
        # Add a section for advanced settings
        st.divider()
        st.subheader("Advanced Settings")
        
        col_left, col_right = st.columns(2)
        
        with col_left:
            # Clear Vector Database button
            if hasattr(lib, 'RAG_AVAILABLE') and lib.RAG_AVAILABLE:
                # Only show this button if RAG is available
                if st.button("Clear Vector Database", type="secondary", use_container_width=True):
                    with st.spinner("Clearing vector database..."):
                        success = lib.clear_vector_database()
                        if success:
                            st.success("Vector database cleared successfully!")
                            # Also clear the vector_db from session state
                            if "vector_db" in st.session_state:
                                del st.session_state["vector_db"]
                            if "selected_doc_sources" in st.session_state:
                                del st.session_state["selected_doc_sources"]
                            if "processed_anonymizations" in st.session_state:
                                del st.session_state["processed_anonymizations"]
                            st.rerun()
                        else:
                            st.error("Failed to clear vector database. If you're seeing a file access error, this could be due to Windows file locks. Try restarting the application.")
                            if st.button("Restart Application", key="restart_app_button"):
                                st.rerun()
            else:
                st.info("RAG functionality not available.")
        
        with col_right:
            # Display current model and provide model information
            st.info(f"Current model: **{st.session_state['openai_model']}**")
            
            # Add model information based on the selected model
            model_info = {
                "gpt-4o": "Latest and most capable model, optimized for performance and cost-effectiveness.",
                "gpt-4-turbo": "Powerful model with strong reasoning capabilities and knowledge up to Apr 2023.",
                "gpt-4": "Original GPT-4 model with high accuracy and reasoning capabilities.",
                "gpt-3.5-turbo": "Fast and cost-effective model for most general tasks. 16K context window."
            }
            
            if st.session_state["openai_model"] in model_info:
                st.caption(model_info[st.session_state["openai_model"]])
            
            # Model rate information
            st.caption("**Note:** Different models have different pricing. Check [OpenAI pricing](https://openai.com/pricing) for details.")
            
            # Add a button to reset API key and change model
            if st.button("Change AI Model", use_container_width=True):
                del st.session_state["gpt_api_key"]
                del st.session_state["openai_model"]
                # Clear vector DB session state if it exists
                if "vector_db" in st.session_state:
                    del st.session_state["vector_db"]
                st.rerun()
            
            # Add a debug toggle (hidden behind a "secret" checkbox)
            if st.checkbox("Enable debug mode", key="debug_toggle", value=False):
                st.session_state.show_debug = True
            else:
                st.session_state.show_debug = False
