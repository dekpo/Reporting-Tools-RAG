# Packages
import streamlit as st
from docx import Document
import pyperclip

# Modules
import io
import json
import base64
from datetime import datetime

# Utils
import lib

# The App config and styling
lib.app_config()

# Sidebar and Navbar
lib.sidebar()

st.title(lib.APP_TITLE)

st.divider()

# Steps
lib.steps(2)

st.divider()

st.header("Reverse Anonymization")

st.markdown("<p>Re-insert entities ( people, organizations, locations...) to your contents using previously saved data.</p>",unsafe_allow_html=True)

st.subheader("Your Content's Entities")
if "saved_gpt_answers" in st.session_state:
    saved_entities = st.session_state["saved_gpt_answers"]["Entities"]
    
    # Check if entities have the expected structure
    if not saved_entities or "Text" not in saved_entities:
        st.warning("The entities data is not in the expected format. This might be because you're accessing a document from a previous session.")
        
        # Create a default structure for entities if missing
        saved_entities = {
            "Text": [],
            "Replacement": [],
            "Category": []
        }
        # Update the session state with the default structure
        st.session_state["saved_gpt_answers"]["Entities"] = saved_entities
    
    st.write(f'You have **{len(saved_entities["Text"])}** entities to reverse.')

    # Add option to manually add entities
    with st.expander("Add New Entity for Reverse Anonymization"):
        col1, col2, col3 = st.columns(3)
        with col1:
            new_text = st.text_input("Original Text (to restore)", key="new_entity_text")
        with col2:
            new_replacement = st.text_input("Anonymized Text (to replace)", key="new_entity_replacement")
        with col3:
            new_category = st.selectbox("Category", ["PERSON", "ORG", "GPE", "DATE", "OTHER"], key="new_entity_category")
        
        if st.button("Add Entity", key="add_entity_btn"):
            if new_text and new_replacement:
                # Add to entities
                saved_entities["Text"].append(new_text)
                saved_entities["Replacement"].append(new_replacement)
                saved_entities["Category"].append(new_category)
                
                # Update session state
                st.session_state["saved_gpt_answers"]["Entities"] = saved_entities
                
                # Clear inputs
                st.session_state["new_entity_text"] = ""
                st.session_state["new_entity_replacement"] = ""
                
                # Show success message
                st.success(f"Added entity: {new_replacement} → {new_text}")
                st.rerun()
            else:
                st.error("Please provide both original text and anonymized text")

    # Display editable entities table
    edited_entities = st.data_editor(saved_entities, disabled=["Replacement","Category"], use_container_width=True)
    
    # Add export/import functionality
    col1, col2, col3 = st.columns(3)
    
    with col1:
        # Export entities as JSON
        if len(saved_entities["Text"]) > 0:
            # Convert entities to JSON
            entities_json = json.dumps(saved_entities, indent=2)
            
            # Create download link
            b64 = base64.b64encode(entities_json.encode()).decode()
            export_filename = f"entities_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
            href = f'<a href="data:application/json;base64,{b64}" download="{export_filename}">Download Entities as JSON</a>'
            st.markdown(href, unsafe_allow_html=True)
    
    with col2:
        # Import entities from JSON
        uploaded_file = st.file_uploader("Import Entities from JSON", type=["json"])
        if uploaded_file is not None:
            try:
                # Read and parse JSON
                entities_data = json.loads(uploaded_file.getvalue().decode())
                
                # Validate structure
                if "Text" in entities_data and "Replacement" in entities_data and "Category" in entities_data:
                    # Update entities
                    saved_entities = entities_data
                    st.session_state["saved_gpt_answers"]["Entities"] = saved_entities
                    
                    # Show success message
                    st.success(f"Imported {len(saved_entities['Text'])} entities successfully!")
                    st.rerun()
                else:
                    st.error("Invalid entities JSON format")
            except Exception as e:
                st.error(f"Error importing entities: {e}")
    
    with col3:
        # Save entities back to the database
        if st.button("Save Entities to Database", key="save_entities_to_db"):
            # Check if we have a document hash
            persist_directory = './chroma_db'
            document_metadata = lib.load_document_metadata(persist_directory)
            
            # Find document hash by title
            doc_title = st.session_state["saved_gpt_answers"]["Title"]
            doc_hash = None
            
            for hash_key, metadata in document_metadata.items():
                if metadata.get('title') == doc_title:
                    doc_hash = hash_key
                    break
            
            if doc_hash:
                # Update entities in the database
                success = lib.update_document_entities(doc_hash, edited_entities)
                if success:
                    st.success(f"Entities saved to database for document '{doc_title}'")
                else:
                    st.error("Failed to save entities to database")
            else:
                st.warning(f"Could not find document '{doc_title}' in the database")
    
    saved_attendees = st.session_state["saved_gpt_answers"]["Attendees"]

    # Check if attendees have the expected structure
    if not saved_attendees or not isinstance(saved_attendees, dict):
        saved_attendees = {}
        st.session_state["saved_gpt_answers"]["Attendees"] = saved_attendees
    elif "Attendee" not in saved_attendees and len(saved_attendees) > 0:
        st.warning("The attendees data is not in the expected format.")
        saved_attendees = {
            "Attendee": [],
            "Replacement": []
        }
        st.session_state["saved_gpt_answers"]["Attendees"] = saved_attendees

    # Initialize edited_attendees to None
    edited_attendees = None
    
    if saved_attendees and "Attendee" in saved_attendees and len(saved_attendees["Attendee"]) > 0:
        st.subheader("The Meeting Attendees")
        st.write(f'You have **{len(saved_attendees["Attendee"])}** attendees to re-insert.')

        edited_attendees = st.data_editor(saved_attendees,disabled=["Replacement"],use_container_width=True)
    
    reverse_now = st.checkbox('**Re-insert All Entities Now ?**',key="reverse_now",value=True,help="This feature can automatically re-insert entities in the content based on your backup.")

    st.divider()
    
    st.write("""<div id='top-content'></div>""",unsafe_allow_html=True)

    st.subheader("Your Anonymized Content Reversed is Below:")

    st.markdown("<p><a href='#download-or-save-your-data'>🔽Download Or Copy Your Data At The Bottom Of This Page🔽</a></p>",unsafe_allow_html=True)

    content = st.container(border=True)

    content.write("""<div class='extracted-content' />""",unsafe_allow_html=True)

    txt = st.session_state["saved_gpt_answers"]["Data"]

    if reverse_now:
        # Check if entities have the expected structure before processing
        if "Replacement" in edited_entities and "Text" in edited_entities and "Category" in edited_entities:
            for term in edited_entities["Replacement"]:
                if term:  # Only process non-empty terms
                    index = edited_entities["Replacement"].index(term)
                    if index < len(edited_entities["Text"]) and index < len(edited_entities["Category"]):
                        replacement = edited_entities["Text"][index]
                        category = edited_entities["Category"][index]
                        txt = txt.replace(term, f'<span class="entity {category.lower()}">{replacement}</span>')
        
        # Check if attendees have the expected structure before processing
        if saved_attendees and "Attendee" in saved_attendees and "Replacement" in saved_attendees:
            for i, attendee in enumerate(saved_attendees["Replacement"]):
                if attendee and i < len(saved_attendees["Attendee"]):
                    replacement = saved_attendees["Attendee"][i]
                    txt = txt.replace(attendee, f'<span class="entity person">{replacement}</span>')
    
    txt = txt.replace("$","\\$")
    txt = txt.replace("\n","<br>")
    content.markdown(txt,unsafe_allow_html=True)
    txt = txt.replace("<br>","\n")
    raw_text = lib.strip_tags(txt)

    st.divider()

    st.subheader("Download Or Save Your Data")
            
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button(f"Copy Your Content To Clipboard",use_container_width=True):
            gpt_answers = raw_text
            pyperclip.copy(gpt_answers)
            st.success("Answers Copied !")
    with col2:
        st.write("OR")
    with col3:
        doc_download = Document()
        doc_title = st.session_state["saved_gpt_answers"]["Title"].replace(">"," ")
        # Filter the title for invalid XML characters
        doc_title = lib.filter_xml_chars(doc_title)
        doc_download.add_heading(doc_title,level=1)
        # Filter out invalid XML characters before adding to the document
        raw_text = lib.filter_xml_chars(raw_text)
        doc_download.add_paragraph(raw_text)
        bio = io.BytesIO()
        doc_download.save(bio)
        if doc_download:
            st.download_button(
            label="**Download Your Content As DOCX File**",
            data=bio.getvalue(),
            type="primary",
            file_name=f"ChatGPT Answers About {doc_title}.docx",
            mime="docx"
            )
    st.divider()
    st.markdown("<p><a href='#top-content'>🔼Go Back To The Top Of This Content🔼</a></p>",unsafe_allow_html=True)
else:
    st.markdown('<p>You have <b>no entity</b> to re-insert. You must first apply anonymization with the previous tool to be able to keep references of your entities.</p>',unsafe_allow_html=True)
    st.page_link(page="pages/01_anonymize.py",label="Got To Anonymize Content",icon=":material/sms:",use_container_width=True)
#st.write(st.session_state)